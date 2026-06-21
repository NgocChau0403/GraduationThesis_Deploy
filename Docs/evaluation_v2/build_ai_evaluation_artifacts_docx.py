from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "Docs" / "evaluation_v2" / "AI_EXPLANATION_EVALUATION_ARTIFACTS.docx"

UCI = ROOT / "Docs" / "evaluation_v2" / "Runs" / "full_208" / "phase11_v3_uci_action_evidence_rerun"
OULAD = ROOT / "Docs" / "evaluation_v2" / "Runs" / "full_208" / "phase12_v3_oulad_action_evidence_rerun"

ROWS = [
    (
        "Báo cáo tổng hợp",
        "UCI report",
        UCI / "UCI_ACTION_EVIDENCE_RERUN_REPORT.md",
        "OULAD report",
        OULAD / "OULAD_ACTION_EVIDENCE_RERUN_REPORT.md",
    ),
    (
        "Báo cáo scoring chi tiết",
        "UCI scoring",
        UCI / "scoring" / "scoring_report__SAMPLE_UCI_POR.md",
        "OULAD scoring",
        OULAD / "scoring" / "scoring_report__SAMPLE_OULAD.md",
    ),
    (
        "So sánh từng task",
        "UCI paired comparison",
        UCI / "scoring" / "aggregates" / "paired_mode_comparison__SAMPLE_UCI_POR.md",
        "OULAD paired comparison",
        OULAD / "scoring" / "aggregates" / "paired_mode_comparison__SAMPLE_OULAD.md",
    ),
    (
        "Log gọi LLM judge",
        "UCI API log",
        UCI / "judge_invocation" / "openai_judge_run_log.jsonl",
        "OULAD API log",
        OULAD / "judge_invocation" / "openai_judge_run_log.jsonl",
    ),
]


def set_run_font(run, size=11, bold=False, color="000000"):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)


def add_hyperlink(paragraph, text, target):
    relationship_id = paragraph.part.relate_to(
        Path(target).resolve().as_uri(),
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    run_fonts = OxmlElement("w:rFonts")
    run_fonts.set(qn("w:ascii"), "Calibri")
    run_fonts.set(qn("w:hAnsi"), "Calibri")
    run_properties.append(run_fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "22")
    run_properties.append(size)
    run.append(run_properties)

    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def build_document():
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("AI Explanation Evaluation Artifacts")
    set_run_font(title_run, size=23, bold=True, color="0B2545")

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(16)
    subtitle_run = subtitle.add_run("UCI and OULAD evaluation reports, scoring outputs, comparisons, and judge logs")
    set_run_font(subtitle_run, size=11, color="555555")

    intro = document.add_paragraph()
    intro.paragraph_format.space_after = Pt(10)
    intro_run = intro.add_run("Use the links below to open the principal AI explanation evaluation artifacts in the project workspace.")
    set_run_font(intro_run)

    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    widths = [2200, 3580, 3580]
    set_table_geometry(table, widths)

    headers = ["Nội dung", "UCI", "OULAD"]
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, "E8EEF5")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        set_run_font(run, bold=True, color="0B2545")

    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_properties.append(repeat)

    for label, uci_text, uci_target, oulad_text, oulad_target in ROWS:
        cells = table.add_row().cells
        for cell in cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)

        label_paragraph = cells[0].paragraphs[0]
        label_paragraph.paragraph_format.space_before = Pt(0)
        label_paragraph.paragraph_format.space_after = Pt(0)
        label_run = label_paragraph.add_run(label)
        set_run_font(label_run, bold=True, color="1F3A5F")

        for cell, text, target in (
            (cells[1], uci_text, uci_target),
            (cells[2], oulad_text, oulad_target),
        ):
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            add_hyperlink(paragraph, text, target)

    set_table_geometry(table, widths)

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(10)
    note.paragraph_format.space_after = Pt(0)
    note_run = note.add_run("Note: Links point to local files in this workspace and require the same project path to remain available.")
    set_run_font(note_run, size=9, color="555555")

    document.core_properties.title = "AI Explanation Evaluation Artifacts"
    document.core_properties.subject = "UCI and OULAD evaluation file index"
    document.core_properties.author = "Graduation Thesis Project"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
