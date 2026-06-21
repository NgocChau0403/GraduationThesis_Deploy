from __future__ import annotations

import json
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BASE_DOCX = ROOT / "thesis" / "THESIS_REPORT_BASE_CLEAN.docx"
OUT_DOCX = ROOT / "thesis" / "THESIS_REPORT_DRAFT_V1.docx"
BUILD_DOCX = ROOT / "thesis" / "THESIS_REPORT_DRAFT_V1.building.docx"

TASK_UCI = ROOT / "Docs/evaluation_logs/task_availability/task_availability_20260619T042411Z_265c90.json"
TASK_OULAD = ROOT / "Docs/evaluation_logs/task_availability/task_availability_20260619T042455Z_0ef6b3.json"
IMPORT_UCI = ROOT / "Docs/evaluation_logs/import_conversion/import_conversion_20260620T080702Z_043fc1.json"
SCORING_UCI = ROOT / "Docs/evaluation_v2/Runs/full_208/phase11_v3_uci_action_evidence_rerun/scoring/scoring_report__SAMPLE_UCI_POR.md"
SCORING_OULAD = ROOT / "Docs/evaluation_v2/Runs/full_208/phase12_v3_oulad_action_evidence_rerun/scoring/scoring_report__SAMPLE_OULAD.md"
PERFORMANCE = ROOT / "Docs/evaluation_logs/system_performance/SYSTEM_PERFORMANCE_EVALUATION.md"


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def load_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def md_value(text: str, label: str, default: str = "") -> str:
    pattern = rf"- {re.escape(label)}:\s*(.+)"
    m = re.search(pattern, text)
    return m.group(1).strip() if m else default


def collect_evidence() -> dict:
    uci_task = load_json(TASK_UCI)
    oulad_task = load_json(TASK_OULAD)
    import_uci = load_json(IMPORT_UCI)
    uci_score = load_text(SCORING_UCI)
    oulad_score = load_text(SCORING_OULAD)

    return {
        "task_availability": {
            "UCI Portuguese": uci_task["summary"],
            "OULAD": oulad_task["summary"],
        },
        "import_conversion": {
            "success": import_uci.get("success"),
            "raw_rows": import_uci.get("raw_rows", {}),
            "canonical_rows": import_uci.get("canonical_rows", {}),
            "fk_errors": import_uci.get("integrity_checks", {}).get("fk_errors"),
            "null_range_errors": import_uci.get("integrity_checks", {}).get("null_range_errors"),
            "feature_checks": import_uci.get("feature_checks", {}),
        },
        "ai_scoring": {
            "UCI Portuguese": {
                "records": md_value(uci_score, "Scored records"),
                "avg_raw": md_value(uci_score, "Average raw weighted score"),
                "avg_final": md_value(uci_score, "Average final score after caps"),
                "verdicts": md_value(uci_score, "Verdict counts"),
                "major_errors": md_value(uci_score, "Highest error severity counts"),
                "avg_delta": md_value(uci_score, "Average delta task-aware minus baseline"),
                "winners": md_value(uci_score, "Winner counts"),
                "pairs": md_value(uci_score, "Pair count"),
                "comparable_pairs": md_value(uci_score, "Comparable pair count"),
            },
            "OULAD": {
                "records": md_value(oulad_score, "Scored records"),
                "avg_raw": md_value(oulad_score, "Average raw weighted score"),
                "avg_final": md_value(oulad_score, "Average final score after caps"),
                "verdicts": md_value(oulad_score, "Verdict counts"),
                "major_errors": md_value(oulad_score, "Highest error severity counts"),
                "avg_delta": md_value(oulad_score, "Average delta task-aware minus baseline"),
                "winners": md_value(oulad_score, "Winner counts"),
                "pairs": md_value(oulad_score, "Pair count"),
                "comparable_pairs": md_value(oulad_score, "Comparable pair count"),
            },
        },
    }


def remove_body_from_heading(doc: Document, heading_text: str) -> None:
    target = None
    for p in doc.paragraphs:
        if p.text.strip().upper() == heading_text.upper():
            target = p._element
            break
    if target is None:
        raise RuntimeError(f"Could not find heading {heading_text!r}")

    body = doc.element.body
    children = list(body)
    start = children.index(target)
    sect_pr = None
    if children and children[-1].tag.endswith("sectPr"):
        sect_pr = children[-1]
    for child in children[start:]:
        if child is sect_pr:
            continue
        body.remove(child)


def set_cell_text(cell, text: str, bold: bool = False, align: WD_ALIGN_PARAGRAPH | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    apply_run_font(run, 13, bold=bold)
    if align:
        p.alignment = align
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_widths(table, widths_in):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)


def insert_table_after(paragraph, rows: int, cols: int):
    container = paragraph._parent
    table = container.add_table(rows=rows, cols=cols, width=Inches(6.3))
    paragraph._p.addnext(table._tbl)
    return table


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_in: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value), align=WD_ALIGN_PARAGRAPH.LEFT if len(str(value)) > 20 else WD_ALIGN_PARAGRAPH.CENTER)
    if widths_in:
        set_table_widths(table, widths_in)
    doc.add_paragraph()
    return table


def apply_run_font(run, size=13, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_paragraph(doc: Document, text: str, style: str | None = None, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if align:
        p.alignment = align
    run = p.add_run(text)
    apply_run_font(run, 13)
    return p


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Paragraph")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        run = p.add_run("- " + item)
        apply_run_font(run, 13)


def add_numbered(doc: Document, items: list[str]):
    for idx, item in enumerate(items, start=1):
        p = doc.add_paragraph(style="List Paragraph")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        run = p.add_run(f"{idx}. " + item)
        apply_run_font(run, 13)


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        apply_run_font(run, 16, bold=True)
        if level <= 2:
            run.font.color.rgb = RGBColor(46, 116, 181)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    return p


def page_break(doc: Document):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_source_note(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    apply_run_font(r, 13, italic=True)
    r.font.color.rgb = RGBColor(90, 90, 90)


def add_reference_entry(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    apply_run_font(r, 13)
    return p


ABBREVIATIONS = [
    ("AI", "Artificial Intelligence"),
    ("API", "Application Programming Interface"),
    ("CPU", "Central Processing Unit"),
    ("CSV", "Comma-Separated Values"),
    ("DBMS", "Database Management System"),
    ("ETL", "Extract, Transform, Load"),
    ("FK", "Foreign Key"),
    ("HTTP", "Hypertext Transfer Protocol"),
    ("JSON", "JavaScript Object Notation"),
    ("JSONL", "JSON Lines"),
    ("KPI", "Key Performance Indicator"),
    ("LA", "Learning Analytics"),
    ("LLM", "Large Language Model"),
    ("LMS", "Learning Management System"),
    ("LRS", "Learning Record Store"),
    ("MVP", "Minimum Viable Product"),
    ("OULAD", "Open University Learning Analytics Dataset"),
    ("PK", "Primary Key"),
    ("RAG", "Retrieval-Augmented Generation"),
    ("RAM", "Random Access Memory"),
    ("REST", "Representational State Transfer"),
    ("RQ", "Research Question"),
    ("SQL", "Structured Query Language"),
    ("UCI", "University of California, Irvine"),
    ("UI", "User Interface"),
    ("URL", "Uniform Resource Locator"),
    ("VLE", "Virtual Learning Environment"),
    ("xAPI", "Experience API"),
]


def write_list_of_abbreviations(doc: Document) -> None:
    """Fill the preserved front-matter abbreviation section in the base template."""
    heading_idx = None
    abstract_idx = None
    body = doc.element.body
    children = list(body)

    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip().upper()
        if text == "LIST OF ABBREVIATIONS":
            heading_idx = i
        elif text == "ABSTRACT" and heading_idx is not None:
            abstract_idx = i
            break

    if heading_idx is None or abstract_idx is None:
        raise RuntimeError("Could not locate LIST OF ABBREVIATIONS section before ABSTRACT")

    heading = doc.paragraphs[heading_idx]
    abstract = doc.paragraphs[abstract_idx]
    heading_el = heading._element
    abstract_el = abstract._element
    start = children.index(heading_el)
    end = children.index(abstract_el)

    # Remove empty placeholder paragraphs or stale abbreviation content, preserving both headings.
    for child in children[start + 1 : end]:
        body.remove(child)

    table = insert_table_after(heading, rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Abbreviation", "Meaning"]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
    for abbreviation, meaning in ABBREVIATIONS:
        row = table.add_row()
        set_cell_text(row.cells[0], abbreviation, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[1], meaning, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_table_widths(table, [1.7, 4.6])

    spacer = doc.add_paragraph()
    table._tbl.addnext(spacer._p)


def normalize_thesis_typography(doc: Document) -> None:
    """Apply the requested thesis type scale while preserving the template cover."""
    for style_name, size, bold in [
        ("Normal", 13, False),
        ("List Paragraph", 13, False),
        ("Heading 1", 16, True),
        ("Heading 2", 16, True),
        ("Heading 3", 16, True),
    ]:
        if style_name not in doc.styles:
            continue
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = bold

    content_started = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.upper() == "ACKNOWLEDGMENTS":
            content_started = True
        style_name = paragraph.style.name if paragraph.style else ""
        is_heading = style_name.startswith("Heading")
        if is_heading:
            for run in paragraph.runs:
                apply_run_font(run, 16, bold=True, italic=bool(run.italic))
        elif content_started:
            for run in paragraph.runs:
                apply_run_font(run, 13, bold=bool(run.bold), italic=bool(run.italic))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        apply_run_font(run, 13, bold=bool(run.bold), italic=bool(run.italic))


def write_abstract(doc: Document):
    add_heading(doc, "ABSTRACT", 1)
    add_paragraph(
        doc,
        "This thesis presents Smart Learning Dashboard with AI-Assisted Data Analysis, "
        "framed as a metadata-driven and evidence-aware learning analytics architecture "
        "for heterogeneous educational data. The motivation is that many educational "
        "dashboards are tightly coupled to one dataset schema, assume that all analyses "
        "remain meaningful after import, and provide AI-generated explanations without "
        "an explicit validity layer connecting the explanation to the available evidence."
    )
    add_paragraph(
        doc,
        "The proposed system separates four concerns that are often mixed together in "
        "dashboard implementations: canonical data representation, analytical task "
        "definition, task availability validation, and AI-assisted interpretation. "
        "Heterogeneous source datasets are mapped into a canonical educational schema. "
        "A task taxonomy is implemented through a metadata-driven task registry. Before "
        "a task is exposed to users, the system validates whether the current dataset "
        "contains the required structural tables, semantic capabilities, analytical "
        "signals, and sufficient evidence. Completed analytical outputs are then passed "
        "to an AI explanation service that is constrained to summarize the available "
        "evidence rather than invent unsupported educational conclusions."
    )
    add_paragraph(
        doc,
        "The thesis evaluates the architecture through import and mapping checks, task "
        "availability logs, AI explanation quality scoring, coverage and human-utility "
        "analysis, and system performance measurements. In the fixed 52-task evaluation "
        "scope, the system identifies 24 executable, 6 partial, and 22 insufficient-data "
        "tasks for the UCI Portuguese dataset, and 44 executable and 8 partial tasks for "
        "OULAD. The AI explanation evaluation shows high overall quality scores while "
        "also revealing that task-aware summarization improves the baseline only modestly "
        "on average and still has cases where the baseline performs better. Performance "
        "benchmarks show that common SQL analytics complete within milliseconds to low "
        "hundreds of milliseconds, while AI explanation latency is dominated by the "
        "external language-model call. A full OULAD import with more than 10 million raw "
        "clickstream rows completes successfully in approximately 17.13 minutes on the "
        "local evaluation machine."
    )
    add_paragraph(
        doc,
        "The main contribution of the thesis is therefore not a new AI model, but a "
        "reusable learning analytics pipeline that makes explicit what data exists, "
        "which analyses are valid, how results are presented, and how AI explanations "
        "remain grounded in verified analytical evidence."
    )


def chapter_1_outline(doc: Document):
    page_break(doc)
    add_heading(doc, "Chapter 1: INTRODUCTION", 1)
    add_heading(doc, "1.1 Research Motivation", 2)
    add_paragraph(
        doc,
        "Educational institutions increasingly collect data from learning management "
        "systems, course records, assessment platforms, attendance logs, and student "
        "information systems. These datasets contain signals about learning behavior, "
        "academic performance, engagement, and risk. In principle, such signals can help "
        "students understand their own progress and help administrators monitor course "
        "quality. In practice, the data is often fragmented across different schemas, "
        "different levels of granularity, and different educational contexts."
    )
    add_paragraph(
        doc,
        "Learning analytics dashboards are a common response to this problem. They convert "
        "raw educational data into charts, indicators, and summaries. However, many dashboards "
        "are built for one fixed dataset or one institutional database. When the data source "
        "changes, the dashboard logic, SQL queries, visualizations, and explanations often "
        "need to be rewritten. This limits reuse and makes it difficult to evaluate the same "
        "analytical task across datasets such as UCI Student Performance and OULAD."
    )
    add_paragraph(
        doc,
        "The recent availability of language models also creates an opportunity to add AI "
        "explanation to dashboards. Yet AI explanation is useful only when it is grounded in "
        "the actual analytical evidence. If the system exposes an analysis that is not valid "
        "for the available dataset, or if the AI layer interprets partial evidence as complete "
        "evidence, the dashboard may produce confident but misleading educational insights."
    )

    add_heading(doc, "1.2 Problem Statement", 2)
    add_paragraph(
        doc,
        "This thesis addresses the following problem: how can a learning analytics system "
        "support heterogeneous educational datasets while ensuring that analytical tasks "
        "are valid for the available data and that AI-generated explanations remain grounded "
        "in verified analytical results?"
    )
    add_paragraph(
        doc,
        "The problem contains three connected subproblems. First, source datasets must be "
        "normalized without erasing important educational meaning. Second, analytical tasks "
        "must be described in a dataset-independent way while still declaring their concrete "
        "data requirements. Third, AI explanations must be connected to completed analytics "
        "outputs, not to arbitrary rows or incomplete context."
    )

    add_heading(doc, "1.3 Research and Practical Gap", 2)
    add_paragraph(
        doc,
        "Prior learning analytics systems often focus on visualization, prediction, or "
        "dashboard usability. These are important directions, but they do not fully solve "
        "the reuse and validity problem. A visualization can be correct for one schema but "
        "invalid for another. A predictive warning can be informative in one dataset but "
        "unsupported in another dataset that lacks temporal engagement or assessment history. "
        "A language-model explanation can sound plausible while silently ignoring missing "
        "fields or insufficient evidence."
    )
    add_paragraph(
        doc,
        "The practical gap is therefore a coordination gap. A dashboard needs a shared data "
        "foundation, a task definition layer, a validity-checking mechanism, and an evidence "
        "boundary for AI. Without this coordination, the system may either hide useful analyses "
        "or expose analyses that should not be trusted."
    )

    add_heading(doc, "1.4 Objectives and Scope", 2)
    add_paragraph(
        doc,
        "The objective of this thesis is to design, implement, and evaluate a thesis-scale "
        "learning analytics dashboard that transforms heterogeneous educational data into "
        "validated analytical tasks and grounded AI explanations."
    )
    add_bullets(
        doc,
        [
            "Design a canonical educational schema that can represent common entities such as students, courses, classes, enrollments, assessments, events, and engagement.",
            "Define an educational task taxonomy and implement it through a metadata-driven task registry.",
            "Validate task availability from actual dataset capabilities before executing analytics.",
            "Generate AI explanations from analytical evidence while limiting unsupported educational claims.",
            "Evaluate the system through task availability, AI explanation quality, coverage and human utility, and performance logs.",
        ],
    )
    add_paragraph(
        doc,
        "The scope is intentionally bounded. The thesis does not propose a new foundation "
        "model, a production-scale institutional platform, or a fully validated early-warning "
        "intervention system. Instead, it focuses on a reusable architecture and an evidence "
        "protocol that make dashboard analytics more explicit and auditable."
    )

    add_heading(doc, "1.5 Research Questions", 2)
    add_bullets(
        doc,
        [
            "RQ1: How can heterogeneous educational datasets be normalized into a reusable analytical structure?",
            "RQ2: How can a task taxonomy determine which analyses are meaningful for each dataset?",
            "RQ3: How effectively can task-aware AI generate evidence-grounded explanations?",
            "RQ4: What do coverage, human-utility, and performance evaluations reveal about the proposed system?",
        ],
    )

    add_heading(doc, "1.6 Main Contributions", 2)
    add_numbered(
        doc,
        [
            "A canonical schema for representing heterogeneous educational datasets in a shared analytical structure.",
            "A task taxonomy operationalized through a metadata-driven task registry and task availability validator.",
            "An evidence-grounded AI explanation pipeline that uses completed analytics outputs as explanation context.",
        ],
    )
    add_paragraph(
        doc,
        "These three technical contributions are assessed through a reproducible evaluation process "
        "covering import and mapping verification, task availability, AI explanation quality, "
        "analytical coverage, and system performance. The evaluation process is supporting evidence "
        "for the contributions rather than a separate architectural contribution."
    )

    add_heading(doc, "1.7 Thesis Structure", 2)
    add_paragraph(
        doc,
        "Chapter 2 reviews background and related work. Chapter 3 describes the proposed "
        "methodology. Chapter 4 explains the implemented system. Chapter 5 reports the "
        "experiments and results. Chapter 6 discusses the findings, limitations, and threats "
        "to validity. Chapter 7 concludes the thesis and outlines future work."
    )


def chapter_1(doc: Document):
    page_break(doc)
    add_heading(doc, "Chapter 1: INTRODUCTION", 1)

    add_heading(doc, "1.1 Research Motivation", 2)
    add_paragraph(
        doc,
        "Educational institutions generate data through student information systems, learning "
        "management systems, assessment platforms, attendance records, course administration, and "
        "online learning environments. These records contain observable evidence about academic "
        "performance, participation, assessment progress, resource use, submission behavior, and "
        "learning context. When interpreted responsibly, such evidence can help students reflect on "
        "their progress and help instructors or administrators monitor cohorts and courses."
    )
    add_paragraph(
        doc,
        "The existence of data does not automatically make it analytically usable. Educational "
        "datasets differ in structure, vocabulary, granularity, and coverage. A compact school-"
        "performance file may combine demographics, family context, absences, and several grades in "
        "one row. An online-learning dataset may distribute comparable concepts across registration, "
        "assessment, submission, course, resource, and clickstream files. Identifiers and score scales "
        "also vary. Consequently, a dashboard built around one source schema cannot be assumed to "
        "operate correctly when the source changes."
    )
    add_paragraph(
        doc,
        "Learning analytics dashboards address part of this challenge by turning educational records "
        "into indicators, charts, tables, and summaries. Their practical value is accessibility: users "
        "should not need to write SQL or inspect thousands of raw rows to understand an assessment "
        "trend or cohort distribution. However, many dashboard implementations bind data ingestion, "
        "queries, chart selection, and interpretation to one fixed database. Adding another dataset "
        "then becomes a software redevelopment task rather than a governed mapping task."
    )
    add_paragraph(
        doc,
        "A second challenge is analytical validity. Successful import only proves that records were "
        "stored; it does not prove that every analytical task is meaningful. A trend requires a time "
        "dimension, a comparison requires adequate groups, an engagement analysis requires recorded "
        "activity, and a relationship analysis requires appropriate non-null variables. If a dashboard "
        "shows every registered task after every import, it can produce empty charts, misleading "
        "aggregates, or confident interpretations of absent evidence."
    )
    add_paragraph(
        doc,
        "The use of large language models adds a third challenge. Natural-language explanation can "
        "make analytical results easier to understand, but fluent text can obscure weak evidence. A "
        "model may overgeneralize from the first rows of a large result, repeat an unsupported number, "
        "infer causality from association, or turn an observable risk indicator into a deterministic "
        "judgment about a student. AI explanation is therefore useful only when it is downstream of "
        "validated analytics and constrained by a traceable evidence packet."
    )
    add_heading(doc, "1.1.1 Thesis Perspective", 3)
    add_paragraph(
        doc,
        "This thesis treats the dashboard as the visible artifact, while the research focus is the "
        "architecture that makes the dashboard reusable and evidence-aware. The system must answer "
        "four questions in sequence: what educational evidence has been imported, what analyses are "
        "defined, which analyses are valid for the active dataset, and what may safely be stated about "
        "a completed result. This sequence motivates a canonical schema, an analytical task taxonomy, "
        "task-availability validation, and evidence-grounded AI explanation."
    )

    add_heading(doc, "1.2 Problem Statement", 2)
    add_paragraph(
        doc,
        "The central research problem is: how can a learning analytics system support heterogeneous "
        "educational datasets while ensuring that analytical tasks are valid for the available data "
        "and that AI-generated explanations remain grounded in verified analytical results?"
    )
    add_paragraph(
        doc,
        "The problem is not solved by implementing a generic CSV upload. Column mapping can make two "
        "files syntactically compatible while leaving their educational meanings different. Nor is it "
        "solved by storing every source in one wide table: such a representation obscures entity grain, "
        "relationships, missing domains, and the provenance of derived values. A reusable analytical "
        "foundation must preserve students, instructional context, enrollment, assessment, results, "
        "events, and engagement as explicit concepts."
    )
    add_paragraph(
        doc,
        "The problem is also not solved by maintaining a list of dashboard pages. An analytical task "
        "must declare its audience, intent, source tables, required capabilities, parameters, query, "
        "output structure, visualization, and explanation strategy. Without such a specification, "
        "data requirements remain hidden in code and cannot be evaluated consistently across datasets."
    )
    add_paragraph(
        doc,
        "Finally, giving a language model direct access to arbitrary database rows does not create a "
        "reliable interpretation layer. The model should receive completed task outputs, deterministic "
        "summaries where necessary, field semantics, known limitations, and safety constraints. The "
        "system must preserve the distinction between observed fields, persisted engineered features, "
        "runtime-derived metrics, and narrative interpretation."
    )
    add_heading(doc, "1.2.1 Connected Subproblems", 3)
    add_table(
        doc,
        ["Subproblem", "Failure if untreated", "Required system response"],
        [
            ["Data heterogeneity", "Queries and dashboards must be rewritten for each source", "Canonical entities, explicit mapping, transformation, and provenance"],
            ["Hidden task assumptions", "Tasks run with missing concepts or inadequate evidence", "Metadata-defined requirements and availability validation"],
            ["Ungrounded interpretation", "Fluent but incomplete, unsupported, causal, or unsafe explanations", "Completed analytics evidence, task-aware summaries, structured output, and guardrails"],
            ["Weak evaluation trace", "Claims depend on demonstrations or screenshots", "Frozen scopes, automated logs, paired comparisons, and performance records"],
        ],
        [1.45, 2.55, 2.5],
    )

    add_heading(doc, "1.3 Research and Practical Gap", 2)
    add_paragraph(
        doc,
        "Research in learning analytics and educational data mining includes dashboards, descriptive "
        "analysis, predictive models, early-warning systems, visual analytics, and AI-assisted "
        "support. These directions demonstrate that educational data can be transformed into useful "
        "signals. Nevertheless, the integration problem addressed here remains important: data "
        "normalization, task definition, task validity, visualization, and AI explanation are often "
        "designed or evaluated as separate concerns."
    )
    add_heading(doc, "1.3.1 Dataset Coupling", 3)
    add_paragraph(
        doc,
        "Many systems assume a stable institutional schema. Their SQL, filters, and chart components "
        "are written against known columns. This approach is efficient for one deployment, but it "
        "limits reuse and makes comparison across educational datasets difficult. Existing ETL or "
        "mapping support may load new records without defining how source concepts correspond to a "
        "shared analytical model. The gap is therefore not file parsing alone, but semantically "
        "controlled conversion into reusable educational entities."
    )
    add_heading(doc, "1.3.2 Task Validity Gap", 3)
    add_paragraph(
        doc,
        "Dashboards commonly hide task requirements inside query code or interface logic. A chart may "
        "fail only after execution, or a task may disappear without explaining which evidence is "
        "missing. A simple supported/unsupported flag also fails to distinguish absent structure, "
        "missing semantic capabilities, analytically weak dimensions, and quantitatively insufficient "
        "data. This thesis addresses the need for an explicit, inspectable bridge between a task "
        "taxonomy and the capabilities of the active dataset."
    )
    add_heading(doc, "1.3.3 AI Grounding and Evaluation Gap", 3)
    add_paragraph(
        doc,
        "AI-generated explanations are frequently evaluated for perceived quality or general "
        "helpfulness, but explanation quality also depends on how analytical evidence is selected and "
        "represented. Passing only the first rows is simple but can omit important groups, extrema, or "
        "temporal patterns. Passing a larger task-aware summary may improve coverage but can also "
        "introduce omissions or prompt complexity. A paired, evidence-aware evaluation is needed to "
        "reveal both advantages and failure cases rather than assume that more elaborate prompting is "
        "always superior."
    )
    add_heading(doc, "1.3.4 Positioning of the Proposed Work", 3)
    add_paragraph(
        doc,
        "The proposed work fills this coordination gap with one traceable pipeline. Source data are "
        "mapped to a shared relational schema. Analytical intents are represented as task metadata. A "
        "four-layer validator checks task availability. Registry-owned SQL produces deterministic "
        "outputs and visualization metadata drives presentation. A separate AI service explains the "
        "completed output within task and safety constraints. The novelty claimed is architectural "
        "integration and explicit evidence handling at thesis scale, not a new database theory, "
        "visualization algorithm, or foundation model."
    )

    add_heading(doc, "1.4 Objectives and Scope", 2)
    add_heading(doc, "1.4.1 General Objective", 3)
    add_paragraph(
        doc,
        "The general objective is to design, implement, and evaluate a metadata-driven and evidence-"
        "aware learning analytics architecture that transforms heterogeneous educational data into "
        "validated analytical tasks and grounded AI explanations."
    )
    add_heading(doc, "1.4.2 Specific Objectives", 3)
    add_numbered(
        doc,
        [
            "Define a canonical educational schema with explicit entity grains and source provenance.",
            "Implement profiling, human-confirmed mapping, transformation, feature calculation, and relational persistence for contrasting datasets.",
            "Define a task taxonomy and operationalize it through a metadata-driven task registry.",
            "Determine task availability from structural, semantic, analytical, and quantitative evidence before ordinary UI execution.",
            "Execute registry-owned parameterized SQL and render results through deterministic visualization metadata and adapters.",
            "Generate structured AI explanations from completed analytics outputs, task context, deterministic summaries, and limitations.",
            "Evaluate import reliability, task availability, AI explanation quality, analytical coverage, operational utility, and system performance.",
        ],
    )
    add_heading(doc, "1.4.3 Research Scope", 3)
    add_paragraph(
        doc,
        "The implemented artifact supports student- and administrator-oriented descriptive and "
        "diagnostic analytics. Its main data cases are UCI Portuguese student performance and OULAD. "
        "The canonical model covers Student, Course, Class, Enrollment, Assessment, AssessmentResult, "
        "Event, and Engagement, together with operational provenance. The fixed primary evaluation "
        "scope contains 52 public analytical tasks."
    )
    add_paragraph(
        doc,
        "The empirical emphasis is Evaluation 3, Evaluation 7, Evaluation 10, and Evaluation 11. "
        "Import, profiling, mapping, SQL, API, and visualization checks provide supporting evidence. "
        "Coverage and utility are interpreted operationally through availability, warnings, "
        "diagnostics, and reviewer inspection; they are not presented as results of a completed large-"
        "scale human-subject study."
    )
    add_heading(doc, "1.4.4 Delimitations", 3)
    add_table(
        doc,
        ["Outside the thesis claim", "Reason for exclusion"],
        [
            ["A novel LLM or predictive model", "The AI contribution concerns grounded interpretation of completed analytics"],
            ["Causal inference", "Observed associations and dashboard patterns do not establish causation"],
            ["A validated early-warning intervention system", "Longitudinal intervention labels and prospective validation are not completed"],
            ["Autonomous high-stakes decisions", "Human review remains necessary for educational interpretation and action"],
            ["Universal production readiness", "Testing is thesis-scale, sequential, local, and limited to two main datasets"],
            ["Proof of human learning improvement", "Evaluation 10 does not constitute a controlled user-impact study"],
        ],
        [3.0, 3.5],
    )

    add_heading(doc, "1.5 Research Questions", 2)
    add_paragraph(
        doc,
        "The study is guided by four research questions. The first three correspond to the technical "
        "contributions. The fourth asks what the supporting empirical evidence reveals about the "
        "complete system; it does not define an additional technical contribution."
    )
    add_table(
        doc,
        ["Research question", "Primary evidence"],
        [
            ["RQ1: How can heterogeneous educational datasets be normalized into a reusable analytical structure?", "Canonical schema, UCI/OULAD mappings, and import/integrity checks"],
            ["RQ2: How can a task taxonomy determine which analyses are meaningful for each dataset?", "Task registry, capability snapshots, and Evaluation 3 availability results"],
            ["RQ3: How effectively can task-aware AI generate evidence-grounded explanations?", "Evaluation 7 paired baseline-versus-task-aware scoring and error analysis"],
            ["RQ4: What do coverage, human-utility, and performance evaluations reveal about the proposed system?", "Evaluations 10 and 11, interpreted with their validity boundaries"],
        ],
        [4.1, 2.4],
    )

    add_heading(doc, "1.6 Main Contributions", 2)
    add_paragraph(
        doc,
        "The thesis makes three connected technical contributions. They should be read as one "
        "pipeline rather than independent features: the canonical schema represents available "
        "evidence, the task taxonomy and validator determine valid analyses, and the AI layer "
        "communicates completed results within an evidence boundary."
    )
    add_heading(doc, "1.6.1 Contribution 1: Canonical Educational Schema", 3)
    add_paragraph(
        doc,
        "The first contribution is an eight-entity canonical relational model for heterogeneous "
        "educational data, accompanied by explicit mapping and feature rules. The schema separates "
        "learner identity and context, instructional organization, enrollment and outcomes, "
        "assessment performance, and learning activity. It supports partial population without "
        "fabricating absent domains and distinguishes raw fields, persisted within-table features, "
        "runtime-derived metrics, and narrative interpretation."
    )
    add_heading(doc, "1.6.2 Contribution 2: Task Taxonomy and Availability Mechanism", 3)
    add_paragraph(
        doc,
        "The second contribution is a task-centered analytical specification. The taxonomy organizes "
        "tasks by audience, scope, analytical intent, data concept, requirements, output shape, and "
        "visualization. The registry operationalizes these concepts through SQL, semantic roles, "
        "output contracts, and explanation metadata. A four-layer validator connects the taxonomy to "
        "the active dataset and returns executable, partial, insufficient-data, or unsupported states "
        "with diagnostic reasons."
    )
    add_heading(doc, "1.6.3 Contribution 3: Evidence-Grounded AI Explanation", 3)
    add_paragraph(
        doc,
        "The third contribution is an explanation pipeline in which the language model is placed after "
        "task validation and SQL execution. The model receives task context, named analytical datasets, "
        "deterministic evidence summaries, known limitations, and safety constraints. Structured "
        "responses retain evidence references, warnings, confidence information, and degradation "
        "status. Baseline and task-aware evidence modes make the effect of evidence preparation "
        "empirically inspectable."
    )
    add_heading(doc, "1.6.4 Evaluation as Supporting Evidence", 3)
    add_paragraph(
        doc,
        "The three contributions are assessed through a reproducible evaluation process. Supporting "
        "checks cover import, profiling, mapping, SQL, API, and visualization behavior. Primary "
        "evidence comes from Evaluation 3 on task availability, Evaluation 7 on AI explanation "
        "quality, Evaluation 10 on coverage and operational utility, and Evaluation 11 on system "
        "performance. The evaluation artifacts strengthen traceability, but the evaluation process is "
        "not claimed as a separate fourth technical contribution."
    )
    add_table(
        doc,
        ["Contribution", "Artifact", "Primary supporting evidence"],
        [
            ["Canonical schema", "Eight canonical entities, mapping rules, and feature taxonomy", "Import conversion and integrity checks; UCI and OULAD persistence"],
            ["Task taxonomy/availability", "Task registry and Layer A-D capability validator", "Evaluation 3 and coverage portion of Evaluation 10"],
            ["Evidence-grounded AI", "Evidence builder, strategies, structured explanation, and safety path", "Evaluation 7 plus explanation review in Evaluation 10"],
        ],
        [1.55, 2.8, 2.15],
    )

    add_heading(doc, "1.7 Thesis Structure", 2)
    add_paragraph(
        doc,
        "Chapter 2 reviews learning analytics, educational dashboards, heterogeneous data sources, "
        "ETL and canonical models, metadata-driven task definitions, task availability, "
        "visualization, and AI explanation. It concludes by locating the thesis within the gap between "
        "fixed dashboards and an evidence-aware analytical architecture."
    )
    add_paragraph(
        doc,
        "Chapter 3 defines the proposed methodology: canonicalization, human-confirmed mapping, task "
        "taxonomy, four-layer availability validation, evidence preparation, AI explanation, and the "
        "evaluation protocol. Chapter 4 describes the concrete implementation across React, Express, "
        "PostgreSQL/Prisma, the task registry, SQL execution, visualization adapters, FastAPI, and "
        "evaluation tooling."
    )
    add_paragraph(
        doc,
        "Chapter 5 reports the experiments and results. It includes supporting technical verification "
        "and emphasizes Evaluations 3, 7, 10, and 11. Chapter 6 interprets the findings in relation to "
        "the research questions and three contributions, compares the architecture with hardcoded "
        "dashboards, and discusses validity threats. Chapter 7 summarizes the conclusions, states the "
        "limitations, and proposes future technical and empirical work."
    )


def chapter_2_outline(doc: Document):
    page_break(doc)
    add_heading(doc, "Chapter 2: BACKGROUND AND RELATED WORK", 1)
    sections = [
        (
            "2.1 Learning Analytics and Educational Data Mining",
            [
                "Learning analytics studies the measurement, collection, analysis, and reporting of data about learners and their contexts. Educational data mining overlaps with learning analytics but often emphasizes computational discovery, prediction, and pattern extraction. Both areas seek to turn educational traces into actionable knowledge.",
                "This thesis uses learning analytics in a pragmatic dashboard sense: the system should provide analyses that help users understand performance, engagement, assessment patterns, and potential risks. The focus is not only on algorithms, but also on whether the system can decide which analyses are valid for the data available.",
            ],
        ),
        (
            "2.2 Educational Dashboards and Data Sources",
            [
                "Educational dashboards summarize complex learning data through visual indicators, charts, and drill-down views. Common data sources include assessment records, enrollment information, demographic attributes, attendance or activity logs, and clickstream events.",
                "A persistent challenge is that datasets differ substantially. UCI Student Performance contains rich demographic and lifestyle variables but limited temporal engagement traces. OULAD contains large-scale virtual learning environment interaction logs and assessment records, but not the same family or lifestyle fields. A reusable dashboard therefore cannot assume that every task is supported by every dataset.",
            ],
        ),
        (
            "2.3 ETL, Schema Mapping, and Canonical Data Models",
            [
                "Extract-transform-load pipelines convert source files into structured database tables. In heterogeneous educational analytics, ETL is not only a technical import step; it is also a semantic mapping problem. Source columns must be interpreted as student attributes, assessment signals, enrollment outcomes, engagement events, or derived features.",
                "A canonical data model provides a stable target schema. It decouples downstream analytics from source-specific file formats. However, a canonical schema must also support partial population because not every dataset contains every educational concept.",
            ],
        ),
        (
            "2.4 Analytical Task Taxonomy and Metadata-Driven Systems",
            [
                "A task taxonomy organizes analyses according to user role, analytical purpose, data requirements, output shape, and visualization type. Instead of hardcoding dashboard pages directly against database tables, a metadata-driven system stores task definitions and uses them to coordinate validation, SQL execution, visualization, and explanation.",
                "In this thesis, the task taxonomy is the conceptual contribution, while the task registry is the implementation mechanism. The registry declares what each task needs and how it should behave when the dataset is missing required evidence.",
            ],
        ),
        (
            "2.5 Task Availability and Data-Sufficiency Validation",
            [
                "Task availability validation answers a simple but critical question: should this analysis be executable for this dataset? A task may be structurally impossible if required tables are absent, semantically unsupported if required capabilities are missing, analytically weak if signals are sparse, or insufficient if the evidence is too limited for a reliable result.",
                "This layer is important because importing a dataset successfully does not mean every analysis is meaningful. The dashboard must distinguish executable, partial, insufficient-data, and unsupported tasks before presenting them to users.",
            ],
        ),
        (
            "2.6 Dashboard Visualization and AI Explanation",
            [
                "Visualization translates query outputs into forms that users can interpret. Different tasks require different visual forms, such as summary tables, trend lines, distribution charts, ranking views, or comparison cards. Visualization must therefore be linked to task metadata rather than chosen arbitrarily.",
                "AI explanation adds narrative interpretation to the dashboard. The main risk is unsupported reasoning. This thesis therefore treats AI explanation as an evidence-bounded interpretation layer. The AI service should explain what the analytics result shows, identify limitations, and avoid claims that cannot be grounded in the provided evidence.",
            ],
        ),
        (
            "2.7 Prior-Work Comparison and Research Gap",
            [
                "The literature and existing dashboard systems show strong progress in analytics, visualization, and AI-assisted educational support. However, many systems remain coupled to fixed schemas, expose analyses without explicit data-sufficiency checks, or evaluate AI explanation without a clear evidence protocol.",
                "The research gap addressed here is the integration of schema normalization, task taxonomy, task validity, and evidence-grounded AI explanation into one coherent learning analytics architecture.",
            ],
        ),
        (
            "2.8 Positioning of This Thesis",
            [
                "This work is positioned as an architecture and evaluation thesis rather than a model-development thesis. Its novelty lies in the way the system coordinates canonical data, task metadata, availability validation, SQL analytics, visualization, and AI explanation.",
                "The dashboard is the user-facing artifact, but the core contribution is the reusable metadata-driven and evidence-aware pipeline behind it.",
            ],
        ),
    ]
    for heading, paras in sections:
        add_heading(doc, heading, 2)
        for para in paras:
            add_paragraph(doc, para)


def chapter_2(doc: Document):
    page_break(doc)
    add_heading(doc, "Chapter 2: BACKGROUND AND RELATED WORK", 1)

    add_heading(doc, "2.1 Learning Analytics and Educational Data Mining", 2)
    add_paragraph(
        doc,
        "Learning analytics concerns the collection, measurement, analysis, and reporting of data "
        "about learners and their contexts for understanding and improving learning. Siemens (2013) "
        "describes the emergence of learning analytics as a discipline shaped by the growth of digital "
        "educational traces and institutional demand for evidence-informed decisions. The field spans "
        "technical analysis, organizational practice, pedagogy, visualization, and ethics rather than "
        "being reducible to one algorithmic technique."
    )
    add_paragraph(
        doc,
        "Educational data mining overlaps with learning analytics but has traditionally emphasized "
        "computational methods for discovering patterns in educational data. Baker and Inventado "
        "(2014) discuss prediction, clustering, relationship mining, discovery with models, and "
        "distillation of data for human judgment. Learning analytics often places stronger emphasis "
        "on stakeholders, interpretation, and institutional action. In practice, the fields share data "
        "sources and methods, and both depend on the quality and meaning of the underlying records."
    )
    add_heading(doc, "2.1.1 Analytical Levels", 3)
    add_paragraph(
        doc,
        "Educational analytics is commonly discussed through descriptive, diagnostic, predictive, "
        "and prescriptive levels. Descriptive analysis summarizes what has been recorded, such as "
        "score distributions or weekly engagement. Diagnostic analysis examines relationships or "
        "contrasts that may help explain an observed pattern. Predictive analysis estimates a future "
        "outcome, while prescriptive analysis recommends actions. The present thesis is intentionally "
        "centered on descriptive and diagnostic tasks. Risk and action outputs are treated as rule-"
        "based indicators grounded in returned evidence, not as validated causal or predictive models."
    )
    add_heading(doc, "2.1.2 Implications for System Design", 3)
    add_paragraph(
        doc,
        "The literature implies that an analytics platform must preserve both computational and human "
        "meaning. A technically correct aggregation can still be inappropriate if its population, time "
        "window, or educational concept is unclear. This thesis therefore treats an analytical task, "
        "rather than a chart, as the governed unit. Each task connects data requirements, execution, "
        "presentation, interpretation, and evaluation."
    )

    add_heading(doc, "2.2 Educational Dashboards and Data Sources", 2)
    add_paragraph(
        doc,
        "Learning dashboards provide visual or textual overviews of learner, course, or institutional "
        "data. Verbert et al. (2013) describe dashboard applications that support awareness and "
        "reflection by presenting educational traces to learners and teachers. Schwendimann et al. "
        "(2017), in a systematic review, show that dashboard research spans multiple stakeholders, "
        "data types, objectives, and visual forms. The review also highlights recurring questions "
        "about evaluation quality and the connection between displayed indicators and learning."
    )
    add_paragraph(
        doc,
        "Student-facing systems deserve separate attention because a technically informative display "
        "may not be understandable or actionable to a learner. Bodily and Verbert (2017) review "
        "student-facing dashboards and recommender systems and identify designs ranging from simple "
        "awareness displays to systems that support comparison, planning, and recommendation. Their "
        "review reinforces the need to distinguish availability of information from demonstrated "
        "educational impact."
    )
    add_heading(doc, "2.2.1 Heterogeneous Educational Sources", 3)
    add_paragraph(
        doc,
        "Dashboard evidence may originate from student information systems, course records, quizzes, "
        "assignment submissions, attendance, discussion forums, learning resources, and fine-grained "
        "platform events. These sources have different grains. A demographic attribute may be stable "
        "per student, an enrollment outcome belongs to a student-class relation, an assessment score "
        "belongs to a student-assessment relation, and a click belongs to an event at a time. Combining "
        "them without explicit entity grains creates duplicate joins and ambiguous metrics."
    )
    add_heading(doc, "2.2.2 UCI Student Performance and OULAD", 3)
    add_paragraph(
        doc,
        "The UCI student-performance data published with the work of Cortez and Silva (2008) represent "
        "a compact school setting with demographic, social, family, study, absence, and grade variables. "
        "The data are useful for performance-oriented analysis but do not contain online engagement "
        "events. The Open University Learning Analytics Dataset, introduced by Kuzilek, Hlosta, and "
        "Zdrahal (2017), contains linked student, registration, course, assessment, and virtual-"
        "learning-environment interaction tables. OULAD therefore supports richer submission and "
        "temporal engagement analysis but lacks several UCI-specific family and lifestyle fields."
    )
    add_table(
        doc,
        ["Dataset property", "UCI Portuguese", "OULAD"],
        [
            ["Typical structure", "One compact student-performance table", "Multiple relational files plus clickstream"],
            ["Performance evidence", "Three grade points and outcomes", "Assessments, submissions, and final results"],
            ["Context evidence", "Demographics, family, lifestyle, support, absences", "Demographics, socioeconomic bands, registration"],
            ["Behavioral evidence", "No VLE clickstream", "Daily resource-level interaction counts"],
            ["Analytical implication", "Strong background/performance tasks; limited online behavior tasks", "Strong temporal/engagement tasks; limited family/lifestyle tasks"],
        ],
        [1.55, 2.45, 2.5],
    )

    add_heading(doc, "2.3 ETL, Schema Mapping, and Canonical Data Models", 2)
    add_paragraph(
        doc,
        "Extract-transform-load pipelines parse source records, transform them to a target model, and "
        "persist validated outputs. In educational analytics, transformation is not only a technical "
        "type-conversion problem. A source field must be interpreted as a student attribute, "
        "instructional context, enrollment value, assessment result, event, or engagement measure. "
        "Incorrect semantic mapping can produce valid database rows and invalid educational analysis."
    )
    add_heading(doc, "2.3.1 Canonicalization versus Source-Specific Modeling", 3)
    add_paragraph(
        doc,
        "Source-specific schemas preserve the original dataset closely and can be efficient for one "
        "application. Their disadvantage is downstream coupling: each query and visualization depends "
        "on source names and joins. A canonical model introduces an intermediate representation. It "
        "requires explicit mapping effort, but downstream tasks can operate on shared entities when "
        "the required concepts are available. The trade-off is therefore additional semantic work at "
        "ingestion in exchange for reuse and clearer provenance later."
    )
    add_paragraph(
        doc,
        "General educational interoperability standards such as Experience API and Caliper Analytics "
        "demonstrate the value of common vocabularies for learning events. However, event exchange is "
        "not identical to an analytics-ready relational model that includes students, classes, "
        "enrollments, assessments, outcomes, and aggregated engagement. The schema in this thesis is a "
        "project-specific analytical canonical model rather than a claim to replace those standards."
    )
    add_heading(doc, "2.3.2 Human-in-the-Loop Mapping", 3)
    add_paragraph(
        doc,
        "Automatic schema matching can use column names, types, value patterns, and known aliases, but "
        "educational meaning can remain ambiguous. A field called result may be a numeric score, a "
        "categorical outcome, or a system status. Human confirmation is therefore an appropriate "
        "control when mappings affect later analysis. Previewing transformed rows shifts validation "
        "earlier in the pipeline and makes mapping decisions inspectable before persistence."
    )
    add_heading(doc, "2.3.3 Raw Fields, Features, Metrics, and Interpretations", 3)
    add_paragraph(
        doc,
        "Analytics systems often use the word feature for several different artifacts. This can blur "
        "the boundary between observed data and interpretation. The present work distinguishes raw "
        "source fields, persisted within-table engineered features, runtime-derived SQL metrics, and "
        "AI-generated interpretations. This separation is consistent with data lineage principles: a "
        "stored pass flag can be recalculated from one score, while a cohort trend or relationship "
        "depends on a task query and should not be stored as an unexplained source fact."
    )

    add_heading(doc, "2.4 Analytical Task Taxonomies and Metadata-Driven Systems", 2)
    add_paragraph(
        doc,
        "A dashboard feature is often described by its visual appearance, such as a line chart or "
        "score card. Visualization research instead separates user intent from visual encoding. "
        "Brehmer and Munzner (2013) propose a multi-level typology that describes why and how a "
        "visualization task is performed and what inputs or outputs are involved. This distinction is "
        "useful for learning analytics because a trend, comparison, distribution, or relationship is "
        "an analytical intent that may have several appropriate visual forms."
    )
    add_heading(doc, "2.4.1 Task as a Declarative Specification", 3)
    add_paragraph(
        doc,
        "Metadata-driven design externalizes behavior that would otherwise be embedded across backend "
        "and frontend code. For this thesis, a task definition includes identity, audience, scope, "
        "analytical concept, source requirements, SQL, parameters, output schema, semantic field roles, "
        "visualization type, and explanation strategy. The registry does not eliminate code: query "
        "execution, validation, adapters, and safety mechanisms remain implemented services. It makes "
        "the variable analytical specification inspectable and versionable."
    )
    add_heading(doc, "2.4.2 Benefits and Risks", 3)
    add_paragraph(
        doc,
        "A registry can improve consistency because multiple consumers read one task definition. It "
        "can improve extensibility because a conventional task does not require a dedicated endpoint "
        "and page. It also creates governance obligations. Metadata can drift from SQL outputs, declare "
        "the wrong chart fields, or omit capability requirements. Registry validation, frozen versions, "
        "output schemas, and runtime diagnostics are therefore necessary companions to metadata-driven "
        "execution."
    )

    add_heading(doc, "2.5 Task Availability and Data-Sufficiency Validation", 2)
    add_paragraph(
        doc,
        "Task availability asks whether a declared analysis is valid for a particular dataset and "
        "context. This question differs from database connectivity and query syntax. A table can exist "
        "while the relevant field is entirely null; a score can exist without enough assessments for "
        "a trend; a comparison can return one group; and an engagement task can be structurally "
        "defined for a dataset that contains no events."
    )
    add_heading(doc, "2.5.1 Levels of Readiness", 3)
    add_table(
        doc,
        ["Readiness level", "Question", "Example failure"],
        [
            ["Structural", "Do required tables and fields exist?", "No engagement table or required score column"],
            ["Semantic", "Do populated fields represent the required educational concept?", "A generic result field does not contain assessment scores"],
            ["Analytical", "Can the intended operation be meaningfully performed?", "One time point for a trend or one group for comparison"],
            ["Sufficiency", "Is the volume and diversity adequate for the declared task?", "Too few students, results, weeks, or positive activity rows"],
        ],
        [1.25, 2.65, 2.6],
    )
    add_paragraph(
        doc,
        "Existing dashboards often handle these conditions with empty-state logic or task-specific "
        "checks. Such controls are useful but difficult to audit when distributed across components. "
        "The gap addressed by this thesis is a reusable capability decision generated from task "
        "metadata and the canonical dataset snapshot. The decision must include reasons and warnings, "
        "not only a Boolean flag."
    )
    add_heading(doc, "2.5.2 Missing Data versus Unsupported Concepts", 3)
    add_paragraph(
        doc,
        "Not all missingness has the same meaning. A few null scores may reduce completeness, while an "
        "entirely absent clickstream means the engagement concept is unsupported. Statistical "
        "imputation may be appropriate for some modeling studies, but unconditional imputation in a "
        "dashboard capability layer can create evidence that the source never contained. The present "
        "approach therefore prefers explicit partial or insufficient states to silent capability "
        "manufacture."
    )

    add_heading(doc, "2.6 Visualization, Feedback, and Human Interpretation", 2)
    add_paragraph(
        doc,
        "Dashboards mediate between analytical output and human interpretation. A visualization should "
        "match the analytical task: temporal progression is commonly shown with position along a time "
        "axis, group comparison with aligned values, and relationships with paired quantitative "
        "coordinates. The chart must also disclose missing or filtered values. Converting null to zero "
        "can turn absence of evidence into an observed low value, which is especially misleading in "
        "student risk or engagement displays."
    )
    add_heading(doc, "2.6.1 Metadata-Driven Rendering", 3)
    add_paragraph(
        doc,
        "A metadata-driven renderer separates analytical semantics from a charting library. Semantic "
        "roles such as category, time, value, series, and tooltip fields are mapped by adapters into "
        "chart-specific structures. This allows shared rendering components while keeping chart choice "
        "deterministic. It also makes missing fields, skipped rows, and null policy available as "
        "diagnostics rather than hiding them inside one component."
    )
    add_heading(doc, "2.6.2 Human Utility Claims", 3)
    add_paragraph(
        doc,
        "Prior dashboard reviews caution against equating availability with impact. A user may see a "
        "chart without understanding it, or understand it without changing behavior. The present thesis "
        "therefore uses a narrow operational meaning of utility: users can discover valid tasks, see "
        "why other tasks are unavailable, inspect evidence limitations, and receive a reviewable "
        "explanation. Claims about learning gains or institutional decision quality require future "
        "human studies."
    )

    add_heading(doc, "2.7 AI Explanation, Grounding, and Evaluation", 2)
    add_paragraph(
        doc,
        "Large language models can translate structured results into readable summaries, but natural "
        "language generation is vulnerable to hallucination. Ji et al. (2023) survey hallucination in "
        "natural language generation and describe failures in which generated content is not supported "
        "by its source. In learning analytics, unsupported statements are especially sensitive because "
        "they can affect perceptions of learner ability, motivation, or risk."
    )
    add_heading(doc, "2.7.1 Evidence Grounding", 3)
    add_paragraph(
        doc,
        "Retrieval-augmented generation demonstrates a broader strategy of supplying external evidence "
        "to a language model rather than relying only on parametric memory (Lewis et al., 2020). The "
        "present system is not a general RAG system: its evidence comes from a completed analytical "
        "task, named result datasets, and deterministic summaries. The shared principle is grounding, "
        "while the implementation is task-output conditioning rather than open-document retrieval."
    )
    add_heading(doc, "2.7.2 Educational Safety and Ethics", 3)
    add_paragraph(
        doc,
        "Learning analytics raises questions about privacy, transparency, agency, and the consequences "
        "of automated interpretation. Slade and Prinsloo (2013) argue that learning analytics creates "
        "ethical dilemmas involving data, identity, responsibility, and institutional power. UNESCO's "
        "guidance on generative AI in education similarly emphasizes human-centered governance and "
        "care with personal data and automated outputs. These concerns motivate restrictions against "
        "unsupported psychological judgments, deterministic predictions, causal claims, and autonomous "
        "high-stakes educational action."
    )
    add_heading(doc, "2.7.3 Evaluating Generated Explanations", 3)
    add_paragraph(
        doc,
        "Explanation evaluation must consider more than fluency. Faithfulness, numerical correctness, "
        "completeness, clarity, relevance, actionability, and safety capture different failure modes. "
        "LLM-as-a-judge methods, including the work of Zheng et al. (2023), offer scalable comparison "
        "but can inherit model and prompt biases. Judge outputs should therefore be schema-validated, "
        "paired on the same evidence, reviewed for severe errors, and interpreted as project-specific "
        "evaluation evidence rather than objective human truth."
    )
    add_heading(doc, "2.7.4 Baseline versus Task-Aware Evidence", 3)
    add_paragraph(
        doc,
        "A first-rows baseline is simple, transparent, and complete when a result contains few rows. It "
        "becomes fragile when ordering is arbitrary or the result is large. Task-aware summarization can "
        "preserve extrema, group totals, temporal changes, distributions, relationships, and action "
        "evidence, but the summary can introduce its own omissions. A paired experiment is therefore "
        "more informative than evaluating only the proposed mode."
    )

    add_heading(doc, "2.8 Prior-Work Comparison and Research Gap", 2)
    add_paragraph(
        doc,
        "The reviewed work provides strong foundations for learning analytics, dashboard design, "
        "educational datasets, visualization tasks, ethical governance, grounding, and automated "
        "evaluation. The unresolved gap is their coordination in one reusable execution path. A fixed "
        "dashboard may visualize one dataset well but not explain validity after another import. A "
        "canonical schema may normalize data without defining analytical intent. A task registry may "
        "execute queries without checking sufficiency. A language model may explain results without a "
        "traceable evidence strategy."
    )
    add_table(
        doc,
        ["Approach", "Typical strength", "Remaining limitation addressed here"],
        [
            ["Learning dashboard", "Accessible monitoring and reflection", "Often coupled to fixed indicators, schemas, or stakeholder contexts"],
            ["Educational dataset/model", "Well-defined source evidence for analysis", "Source-specific structure does not create cross-dataset task reuse"],
            ["ETL/canonical model", "Consistent target representation", "Does not by itself determine which analyses are meaningful"],
            ["Visualization task taxonomy", "Separates intent from visual encoding", "Does not specify educational data readiness or SQL execution"],
            ["LLM explanation", "Readable interpretation and summarization", "Can omit evidence or introduce unsupported claims"],
            ["Proposed architecture", "Connects schema, task intent, availability, execution, visualization, and grounded explanation", "Evaluated at thesis scale; not yet institutionally deployed"],
        ],
        [1.65, 2.25, 2.6],
    )
    add_heading(doc, "2.8.1 Positioning of This Thesis", 3)
    add_paragraph(
        doc,
        "The thesis is positioned as a metadata-driven and evidence-aware learning analytics "
        "architecture. Its three technical contributions are a canonical educational schema, a task "
        "taxonomy combined with task-availability validation, and an evidence-grounded AI explanation "
        "pipeline. The evaluation protocol supplies supporting evidence, with emphasis on task "
        "availability, explanation quality, coverage and operational utility, and performance."
    )
    add_paragraph(
        doc,
        "The contribution is not any individual technique in isolation. Relational modeling, task "
        "typologies, SQL analytics, chart adapters, language models, and automated scoring all exist in "
        "prior work. The contribution lies in integrating them around explicit boundaries: source data "
        "is separated from canonical evidence, task intent from dataset capability, deterministic "
        "analysis from natural-language interpretation, and observed evaluation evidence from broader "
        "claims that remain untested."
    )
    add_heading(doc, "2.9 Chapter Summary", 2)
    add_paragraph(
        doc,
        "This chapter reviewed the foundations required by the proposed system. Learning analytics and "
        "educational data mining explain the analytical context; dashboard studies motivate accessible "
        "presentation while cautioning against weak impact claims; UCI and OULAD illustrate source "
        "heterogeneity; canonicalization and metadata support reuse; task readiness protects validity; "
        "and grounding, safety, and paired evaluation constrain AI explanation. Chapter 3 translates "
        "these foundations into the proposed methodology."
    )


def chapter_3(doc: Document):
    page_break(doc)
    add_heading(doc, "Chapter 3: PROPOSED METHODOLOGY", 1)

    add_heading(doc, "3.1 Overall Research Approach", 2)
    add_paragraph(
        doc,
        "This research follows a design-and-evaluation methodology. The research artifact is "
        "a working learning analytics platform, but the object being investigated is the "
        "architecture behind that platform: how heterogeneous educational data can be "
        "normalized, connected to explicitly defined analytical tasks, checked for task-level "
        "validity, and translated into natural-language explanations without detaching the "
        "explanation from its evidence. The methodology therefore combines system design, "
        "implementation, controlled execution on two datasets, and empirical evaluation of the "
        "resulting analytical behavior."
    )
    add_paragraph(
        doc,
        "The central unit of analysis is an analytical task rather than a dashboard page or an "
        "individual chart. A task represents a governed analytical intent, such as describing a "
        "student's assessment trajectory, comparing groups, examining engagement over time, or "
        "summarizing a cohort-level risk signal. Treating the task as the unit of analysis allows "
        "the same definition to govern data requirements, SQL execution, output structure, "
        "visualization, explanation strategy, and evaluation. This choice is essential because a "
        "chart can be rendered successfully even when its underlying data are incomplete or "
        "semantically unsuitable. The task definition makes those assumptions explicit."
    )
    add_paragraph(
        doc,
        "The proposed method decomposes the end-to-end process into four coordinated layers. "
        "First, a canonical educational schema provides a common representation for source data. "
        "Second, a task taxonomy and registry describe the analyses the system intends to support. "
        "Third, a capability-based validator determines whether each task is valid for the "
        "currently imported dataset. Fourth, an evidence-aware explanation process converts "
        "completed analytical outputs into constrained narrative interpretations. These layers "
        "are evaluated through reproducible runners and logs rather than through interface "
        "demonstration alone."
    )
    add_table(
        doc,
        ["Methodological layer", "Research question addressed", "Primary artifact"],
        [
            ["Canonicalization", "How can heterogeneous educational sources be represented through one analytical foundation?", "Eight canonical domain tables and persisted features"],
            ["Task specification", "How can analytical behavior be defined independently of a specific dashboard or source file?", "Taxonomy and metadata-driven task registry"],
            ["Availability validation", "How can the system prevent execution when required educational evidence is absent or insufficient?", "Four-layer capability decision and diagnostic output"],
            ["Evidence-aware explanation", "How can AI explanations remain tied to completed analytical results and task intent?", "Structured evidence packet, constrained prompt, and validated response"],
            ["Evaluation", "How reliably and efficiently does the complete pipeline behave across contrasting datasets?", "Ground truth, paired comparisons, review records, and performance logs"],
        ],
        [1.45, 3.15, 1.9],
    )
    add_heading(doc, "3.1.1 Methodological Principles", 3)
    add_paragraph(
        doc,
        "Five principles guide the design. The first is separation of concerns: imported fields, "
        "persisted engineered features, runtime-derived metrics, and AI interpretations are kept "
        "conceptually distinct. The second is metadata-driven behavior: analytical intent and "
        "presentation requirements are declared in a registry instead of being scattered across "
        "service code and interface components. The third is validity before execution: the "
        "presence of a task in the registry does not imply that the task is suitable for every "
        "dataset. The fourth is evidence before narrative: the language model receives completed "
        "query results and deterministic summaries rather than direct unrestricted database "
        "access. The fifth is reproducibility: major decisions and outputs are persisted in "
        "machine-readable logs so that reported results can be reconstructed."
    )
    add_heading(doc, "3.1.2 End-to-End Research Workflow", 3)
    add_numbered(
        doc,
        [
            "Profile the uploaded source files and identify their structural and semantic characteristics.",
            "Map source columns and records into the canonical educational schema, with human confirmation when inference is uncertain.",
            "Create a dataset capability snapshot from populated canonical fields, engineered features, dimensions, and record counts.",
            "Load the fixed task registry and evaluate every task against the capability snapshot and task-specific thresholds.",
            "Execute only tasks permitted by the availability policy and normalize their analytical outputs into a stable response contract.",
            "Prepare an evidence packet from the task metadata, completed query output, deterministic summaries, warnings, and audience context.",
            "Generate and validate the AI explanation, then record outputs, errors, latency, token usage, and evaluation artifacts.",
        ],
    )
    add_paragraph(
        doc,
        "This workflow establishes a traceable chain from source evidence to user-facing insight. "
        "A claim in an explanation can be traced backwards to an evidence item, a query result, a "
        "task definition, and ultimately a set of canonical records. Conversely, missing source "
        "evidence propagates forward as a capability limitation, a warning, or a disabled task."
    )

    add_heading(doc, "3.2 Canonical Educational Data Schema", 2)
    add_paragraph(
        doc,
        "The canonical schema addresses structural heterogeneity by translating source-specific "
        "files into a shared relational model. The method does not assume that all educational "
        "datasets contain the same variables. Instead, it defines a stable set of domain entities "
        "and permits dataset-specific fields to be absent when the source does not contain the "
        "corresponding concept. This makes normalization possible without fabricating values."
    )
    add_heading(doc, "3.2.1 Design Principles and Entity Grain", 3)
    add_paragraph(
        doc,
        "Each table is defined by an explicit grain. The grain specifies what one row represents "
        "and prevents accidental duplication during joins. Student describes one learner within a "
        "source dataset; course describes a module; class describes a specific course presentation "
        "or run; enrollment links one learner to one class; assessment describes one graded "
        "activity; assessment_result records one enrollment's result for one assessment; event "
        "describes a learning resource or activity type; and engagement records one enrollment's "
        "interaction with one event on one relative day. Surrogate or composite identifiers retain "
        "source-dataset context so that records from different imports do not collide."
    )
    add_table(
        doc,
        ["Canonical table", "Row grain", "Analytical role"],
        [
            ["student", "One student x source dataset", "Identity, demographics, socioeconomic context, and support signals"],
            ["course", "One course or module", "Stable instructional subject or module identity"],
            ["class", "One course x presentation/run", "Temporal and organizational teaching context"],
            ["enrollment", "One student x class", "Registration, studied credits, withdrawal, and final outcome"],
            ["assessment", "One graded activity", "Assessment type, weight, and due-time context"],
            ["assessment_result", "One enrollment x assessment", "Score, submission status, and performance evidence"],
            ["event", "One learning resource/activity", "Resource type and planned availability"],
            ["engagement", "One enrollment x event x day", "Temporal interaction counts and engagement evidence"],
        ],
        [1.35, 2.05, 3.1],
    )
    add_paragraph(
        doc,
        "The relational links create three main analytical paths. The performance path joins "
        "student, enrollment, class, assessment, and assessment_result. The activity path joins "
        "student, enrollment, class, event, and engagement. The contextual path connects learner "
        "attributes and support variables to enrollment outcomes or performance measures. Tasks "
        "may use one or more paths, but the task registry must declare the tables and fields it "
        "requires."
    )
    add_heading(doc, "3.2.2 Cross-Dataset Normalization", 3)
    add_paragraph(
        doc,
        "Normalization is semantic as well as structural. Equivalent source values are converted "
        "to a common representation where the equivalence is defensible. Examples include "
        "normalizing binary gender codes, converting source outcomes into a shared final-result "
        "vocabulary, translating residence codes into urban, rural, or unknown categories, and "
        "representing assessment scores on a common 0-100 scale. Source values are preserved or "
        "marked unknown when no valid correspondence exists. Dataset-specific concepts remain "
        "nullable rather than being imputed solely to satisfy the schema."
    )
    add_paragraph(
        doc,
        "The two evaluation datasets deliberately stress different parts of the schema. The UCI "
        "student-performance data provide rich demographic, family, lifestyle, support, absence, "
        "and grade variables, but do not provide longitudinal virtual-learning-environment "
        "clickstream records. OULAD provides course presentations, registrations, assessments, "
        "submissions, and large-scale daily engagement events, while some UCI-specific background "
        "attributes are unavailable. The canonical schema preserves these differences and allows "
        "the availability validator to determine their analytical consequences."
    )
    add_heading(doc, "3.2.3 Field and Feature Taxonomy", 3)
    add_paragraph(
        doc,
        "To avoid conflating transformation with analysis, the method distinguishes four types of "
        "representation. A raw field is imported from the source with only type or vocabulary "
        "normalization. An engineered feature is computed from fields in the same canonical row "
        "and persisted during transformation. A derived metric is calculated at query runtime, "
        "often through aggregation or joins. An analytical interpretation is a natural-language "
        "description generated after the metric exists. This distinction defines where each "
        "operation is allowed and what provenance must be retained."
    )
    add_table(
        doc,
        ["Representation", "Storage/time", "Examples", "Methodological constraint"],
        [
            ["Raw field", "Canonical column after import", "score, final_result, engagement_count", "Must correspond to an observed source value"],
            ["Engineered feature", "Persisted at ETL time", "score_normalized, pass_flag, log_click_score", "Computed from fields within the same table row"],
            ["Derived metric", "SQL result at task runtime", "average score, pass rate, total clicks, trend slope", "Query and aggregation logic remain explicit"],
            ["Interpretation", "AI response after execution", "A cautious description of a decline or group difference", "Must not be treated as stored fact or causal proof"],
        ],
        [1.25, 1.35, 2.15, 1.75],
    )
    add_paragraph(
        doc,
        "Only within-table features are persisted in the current methodology. For example, "
        "pass_flag can be computed from a normalized score in assessment_result, and "
        "log_click_score can be computed from engagement_count in engagement. A value such as "
        "submission delay requires joining a result to its assessment due date and is therefore "
        "computed as a runtime metric. This rule keeps ETL deterministic and prevents hidden "
        "cross-table analytical assumptions from becoming database facts."
    )
    add_heading(doc, "3.2.4 Partial Coverage and Missing Data", 3)
    add_paragraph(
        doc,
        "Partial population of the canonical schema is an expected condition, not automatically "
        "an import failure. Missingness is handled at three levels. Structural absence means a "
        "source does not contain a domain at all, such as engagement events in UCI. Attribute "
        "absence means that the entity exists but a particular field is unavailable. Record-level "
        "missingness means that a field exists but is null for some records. Structural and "
        "attribute absence are represented explicitly in the dataset capability snapshot, while "
        "record-level completeness can be included in analytical suitability and sufficiency "
        "checks. The method avoids unconditional statistical imputation because invented values "
        "could make unsupported tasks appear valid."
    )

    add_heading(doc, "3.3 Import and Mapping Method", 2)
    add_paragraph(
        doc,
        "The import methodology converts one or more source files into canonical entities while "
        "retaining enough diagnostics to review the conversion. It combines automated profiling "
        "and suggestion with human confirmation. This human-in-the-loop design is used because "
        "column-name similarity alone cannot establish educational meaning reliably; for example, "
        "a field named result may represent a raw score, a categorical outcome, or an execution "
        "status depending on the source."
    )
    add_heading(doc, "3.3.1 Source Profiling and Dataset Detection", 3)
    add_paragraph(
        doc,
        "Profiling begins before database insertion. For each file, the profiler determines the "
        "delimiter and encoding, reads the header, samples values, estimates data types, counts "
        "rows and missing values, and identifies candidate key or categorical fields. For known "
        "datasets, characteristic file names and column signatures are used to suggest the source "
        "family and entity role. Detection is advisory: it accelerates configuration but does not "
        "replace validation of the proposed mappings."
    )
    add_paragraph(
        doc,
        "When several related files are uploaded, profiling also identifies relationships among "
        "them. OULAD, for example, separates student information, registration, assessment, result, "
        "resource, and clickstream data. These files must be associated before canonical foreign "
        "keys can be constructed. In contrast, the UCI student-performance file contains many "
        "learner, context, and grade fields in a single row and must be decomposed into canonical "
        "entities. The profiler therefore supplies evidence for two different transformation "
        "patterns: integration of multiple normalized source files and decomposition of a flat "
        "source table."
    )
    add_heading(doc, "3.3.2 Mapping and Human Confirmation", 3)
    add_paragraph(
        doc,
        "A mapping specification links a source field to a canonical entity, canonical attribute, "
        "and optional transformation. Suggestions may use normalized name similarity, compatible "
        "source and target types, observed value patterns, known aliases, and previously confirmed "
        "mappings. High confidence reduces manual effort, but every mapping remains inspectable. "
        "Low-confidence or ambiguous mappings can be corrected before the import is committed."
    )
    add_table(
        doc,
        ["Mapping component", "Purpose", "Example"],
        [
            ["Entity assignment", "Select the canonical table represented by a source field or file", "studentInfo.csv to student/enrollment context"],
            ["Field assignment", "Select the canonical attribute", "sex to student.gender"],
            ["Value transform", "Normalize vocabulary, type, or scale", "M/F to male/female; U/R to urban/rural"],
            ["Identifier rule", "Construct stable primary and foreign keys", "student + module + presentation to enrollment_id"],
            ["Null policy", "Preserve unavailable concepts without fabricating evidence", "No UCI clickstream to null engagement capability"],
        ],
        [1.45, 3.05, 2.0],
    )
    add_paragraph(
        doc,
        "A transformed-row preview is produced before persistence when the mapping interface is "
        "used. The preview exposes representative canonical rows after transformations, allowing "
        "incorrect types, unexpected categorical mappings, malformed identifiers, and misplaced "
        "entity assignments to be detected earlier. Confirmed manual aliases may be retained as "
        "mapping memory for future imports, but they are still constrained by target schema and "
        "type validation."
    )
    add_heading(doc, "3.3.3 Transformation, Feature Engineering, and Persistence", 3)
    add_numbered(
        doc,
        [
            "Parse and type-cast source values according to the confirmed mapping.",
            "Normalize categorical vocabularies and score scales without changing the source meaning.",
            "Generate canonical identifiers and resolve foreign-key relationships.",
            "Compute permitted within-table engineered features and retain nulls when prerequisites are unavailable.",
            "Insert entities in dependency order and batch large fact records to control memory and transaction size.",
            "Verify row counts, key integrity, required ranges, and engineered-feature formulas, then write an import log.",
        ],
    )
    add_paragraph(
        doc,
        "Persistence follows relational dependency order: core course and class context is created "
        "before enrollment; students must exist before their enrollments; assessments and events "
        "must exist before result and engagement facts that reference them. Large engagement "
        "sources are processed in batches. A failed integrity check is recorded as an import "
        "failure or warning rather than silently discarding inconsistent records."
    )
    add_heading(doc, "3.3.4 Import Quality Criteria", 3)
    add_paragraph(
        doc,
        "An import is considered technically successful only when the pipeline completes and its "
        "post-import checks satisfy the expected constraints. The checks include source-to-target "
        "row reconciliation where a one-to-one relation is expected, uniqueness of canonical keys, "
        "absence of orphan foreign keys, score and categorical range validity, and direct "
        "recalculation of sampled engineered features. Because some source records are legitimately "
        "aggregated or decomposed, row reconciliation is defined per mapping rather than as one "
        "global equality. The import log records raw and canonical row counts so that such changes "
        "remain explainable."
    )

    add_heading(doc, "3.4 Educational Analytics Task Taxonomy", 2)
    add_paragraph(
        doc,
        "The task taxonomy formalizes the analytical space that the platform intends to support. "
        "Its purpose is not merely to group interface features; it provides a vocabulary for "
        "specifying analytical intent, audience, evidence requirements, execution behavior, and "
        "interpretation strategy. The taxonomy separates the conceptual definition of a task from "
        "the source dataset that happens to make it executable."
    )
    add_heading(doc, "3.4.1 Taxonomy Dimensions", 3)
    add_paragraph(
        doc,
        "Tasks are classified along multiple dimensions because no single category adequately "
        "describes analytical behavior. Role indicates who will interpret the result. Scope "
        "indicates whether the unit is one student, a comparison, a cohort, or an instructor-level "
        "view. Analysis type describes the operation, while data concept states the educational "
        "phenomenon being examined. Required capabilities state the minimum evidence, and optional "
        "capabilities identify enrichments that can improve the result without being mandatory. "
        "Output and visualization metadata define a deterministic presentation contract."
    )
    add_table(
        doc,
        ["Dimension", "Representative values", "Methodological function"],
        [
            ["Audience/role", "Student; administrator or instructor", "Controls language, aggregation level, and privacy expectations"],
            ["Scope", "Single student; comparison; cohort; instructor", "Defines the analytical unit and required filters"],
            ["Intent/type", "Summary; trend; comparison; distribution; correlation; behavioral; risk", "Selects analytical and explanation strategy"],
            ["Data concept", "Performance; engagement; submission; outcome; support; risk", "Connects technical fields to educational meaning"],
            ["Capability", "Scores; outcomes; demographics; time; engagement; group diversity", "Supplies task-availability requirements"],
            ["Output shape", "Metric; grouped rows; time series; ranking; relationship; action evidence", "Defines the response schema"],
            ["Visualization", "Card; bar; line; histogram; scatter; heatmap; table", "Selects deterministic rendering behavior"],
        ],
        [1.25, 2.45, 2.8],
    )
    add_heading(doc, "3.4.2 Student and Administrative Analytical Scopes", 3)
    add_paragraph(
        doc,
        "Student-facing tasks use an individual enrollment as their principal context and are "
        "designed to support reflection on observable performance, activity, punctuality, or "
        "available support indicators. Administrative tasks aggregate across students, classes, "
        "courses, demographic groups, assessment types, or time periods. They are intended for "
        "monitoring and comparison rather than for automatic high-stakes decisions. A task can "
        "share an analytical method across roles while using different aggregation and explanation "
        "rules. For example, a trend task may describe one student's assessment trajectory or an "
        "entire cohort's weekly engagement."
    )
    add_heading(doc, "3.4.3 Operationalization through the Task Registry", 3)
    add_paragraph(
        doc,
        "The taxonomy is operationalized through a fixed registry containing machine-readable "
        "task definitions. A registry entry declares the task identifier and name; role and scope; "
        "description and data concept; source tables and key fields; analysis type; parameterized "
        "SQL; query labels; visualization configuration and semantic field roles; analysis context; "
        "explanation strategy and target audience; and dataset-compatibility requirements. The "
        "registry is therefore both an analytical specification and a governance artifact."
    )
    add_table(
        doc,
        ["Registry field group", "Examples", "Consumer"],
        [
            ["Identity and purpose", "task_id, task_name, description, roles, scope", "Task browser and audit logs"],
            ["Execution", "sql_query, query_labels, parameters, source_tables", "SQL execution service"],
            ["Validity", "required fields/features, analytical conditions, minimum counts", "Capability validator"],
            ["Presentation", "viz_type, semantic_roles, axes, labels", "Visualization adapter"],
            ["Explanation", "analysis_context, explanation_strategy, target_audience", "Evidence builder and AI service"],
        ],
        [1.55, 2.9, 2.05],
    )
    add_paragraph(
        doc,
        "This declarative design reduces coupling. Adding or revising a conventional task usually "
        "requires changing one governed metadata record rather than adding a dedicated API route, "
        "chart component, and prompt. It also improves comparability in evaluation: the same task "
        "scope is applied to both datasets, and differences in availability can be attributed to "
        "dataset capabilities rather than to different feature lists."
    )
    add_heading(doc, "3.4.4 Deterministic Execution and Presentation", 3)
    add_paragraph(
        doc,
        "The task registry constrains both execution and rendering. SQL templates are parameterized "
        "to prevent user filters from altering query structure. Query labels normalize one or more "
        "result sets into named datasets. Visualization metadata assigns semantic roles such as "
        "category, value, time, series, and tooltip fields. Consequently, the AI model does not "
        "choose the chart type and the frontend does not infer analytical meaning from arbitrary "
        "column order. This deterministic boundary is important for evaluating explanation quality "
        "independently of visualization selection."
    )

    add_heading(doc, "3.5 Task Availability Method", 2)
    add_paragraph(
        doc,
        "Task availability is the decision process that connects an intended analysis to the "
        "evidence actually present in a dataset. Formally, for a task t and dataset capability "
        "profile D, the validator evaluates a sequence of requirement predicates covering "
        "structure, semantic signals, analytical suitability, and quantitative sufficiency. The "
        "output is not only a Boolean executable flag; it includes a status, layer-level outcomes, "
        "warnings, and missing requirements so that the decision is explainable."
    )
    add_heading(doc, "3.5.1 Dataset Capability Snapshot", 3)
    add_paragraph(
        doc,
        "After import, the system constructs a capability snapshot from the canonical database. "
        "The snapshot records populated tables and fields, available engineered features, known "
        "semantic dimensions, distinct group counts, temporal coverage, and record counts relevant "
        "to common sufficiency rules. Capabilities are based on stored evidence rather than on the "
        "declared dataset name. Therefore, a custom dataset can support a task if it populates the "
        "required canonical concepts, while a known dataset can fail a task if its import is "
        "incomplete."
    )
    add_heading(doc, "3.5.2 Four Validation Layers", 3)
    add_table(
        doc,
        ["Layer", "Question", "Representative checks", "Failure meaning"],
        [
            ["A. Structural", "Can the query be formed?", "Required canonical tables and columns exist", "Physical requirements are absent"],
            ["B. Semantic", "Does the stored data represent the intended concept?", "Required normalized or engineered signals are populated", "The field shape exists but the educational signal does not"],
            ["C. Analytical", "Can the intended operation be meaningfully performed?", "Time points for trends, multiple groups for comparison, two variables for correlation", "The operation is structurally possible but analytically weak"],
            ["D. Sufficiency", "Is there enough evidence?", "Minimum students, records, assessments, temporal points, or non-null coverage", "Evidence volume falls below a task-specific threshold"],
        ],
        [1.05, 1.55, 2.7, 1.2],
    )
    add_paragraph(
        doc,
        "The layers are evaluated in order because later checks depend on earlier ones. Data-count "
        "queries are not meaningful when the required table is absent, and an analytical trend "
        "check is not meaningful when no temporal field is populated. Early failure nevertheless "
        "produces a diagnostic record rather than a generic error. Where safe, independent checks "
        "may continue so that the user receives a complete list of missing requirements."
    )
    add_heading(doc, "3.5.3 Status Decision Policy", 3)
    add_paragraph(
        doc,
        "Four statuses communicate the result. Executable means that all mandatory requirements "
        "pass. Partial means that the core task can run but one or more optional signals or quality "
        "conditions are absent; the output must be accompanied by warnings and a reduced "
        "interpretation boundary. Insufficient_data means that the required concept exists but the "
        "available records do not meet the declared quantitative threshold. The current prototype "
        "may still execute this state with explicit warnings for diagnostic and degraded-output use. "
        "Unsupported means that mandatory structural requirements are absent, so the registered "
        "query cannot be formed safely."
    )
    add_table(
        doc,
        ["Status", "Execution policy", "User-facing behavior"],
        [
            ["executable", "Permit normal execution", "Show task and result without capability warning"],
            ["partial", "Permit only the defined degraded path", "Show limitations and omit unsupported claims or fields"],
            ["insufficient_data", "Permit the controlled soft-failure path", "Show unmet thresholds, warnings, and limited evidence"],
            ["unsupported", "Block execution", "Identify missing tables, fields, or semantic capabilities"],
        ],
        [1.35, 2.25, 2.9],
    )
    add_paragraph(
        doc,
        "The decision can be summarized as follows. If a mandatory structural predicate fails, "
        "status is unsupported. Otherwise, if a mandatory quantitative sufficiency predicate "
        "fails, status is insufficient_data. If semantic capabilities or non-blocking analytical "
        "conditions fail, status is partial. "
        "Only when mandatory predicates pass without degradation is status executable. Task-level "
        "metadata may refine this policy when a safe degraded result has been explicitly defined."
    )
    add_heading(doc, "3.5.4 Ground-Truth Construction for Availability", 3)
    add_paragraph(
        doc,
        "Availability evaluation requires an expected status for every dataset-task pair. Ground "
        "truth is constructed from three sources: the task's declared requirements, the documented "
        "content of the dataset, and direct inspection of the imported canonical database. The "
        "expected matrix is reviewed to distinguish absent capabilities from low record counts and "
        "to ensure that optional requirements do not incorrectly become blockers. The automated "
        "runner then compares actual and expected status, executable flag, warnings, and missing "
        "requirements. This produces more diagnostic evidence than accuracy on a Boolean flag alone."
    )

    add_heading(doc, "3.6 AI Explanation Method", 2)
    add_paragraph(
        doc,
        "The AI explanation method treats the language model as an interpretation layer, not as an "
        "independent analytics engine. The model does not select database tables, generate hidden "
        "queries, or decide whether the task is valid. Explanation is requested only after task "
        "availability has been resolved and the parameterized analytical query has completed. This "
        "ordering establishes an evidence boundary: the model may organize and communicate supplied "
        "findings, but it must not invent unavailable measurements."
    )
    add_heading(doc, "3.6.1 Evidence Preparation", 3)
    add_paragraph(
        doc,
        "The evidence builder combines task metadata with analytical output. The packet includes the "
        "task identifier, purpose, role, analytical scope, analysis context, explanation strategy, "
        "query labels, field semantics, result-row count, selected or summarized results, applicable "
        "warnings, and known limitations. Numerical summaries are computed deterministically before "
        "prompt construction when the task strategy requires them. Evidence items receive stable "
        "identifiers so that generated insights can reference their basis."
    )
    add_table(
        doc,
        ["Evidence component", "Function in the prompt", "Provenance"],
        [
            ["Task metadata", "Defines intent, audience, and required explanation behavior", "Versioned task registry"],
            ["Query schema and labels", "Defines the meaning of fields and result sets", "SQL execution contract"],
            ["Primary evidence", "Provides observed result rows or aggregate values", "Completed parameterized query"],
            ["Derived evidence", "Provides deterministic totals, rates, ranges, rankings, or trend descriptors", "Evidence summarizer with source links"],
            ["Limitations and warnings", "Constrains confidence and unsupported interpretation", "Availability and data-quality checks"],
            ["Action evidence", "Supplies actions or recommendations already produced by deterministic rules", "Returned query fields or action checker"],
        ],
        [1.55, 2.8, 2.15],
    )
    add_heading(doc, "3.6.2 Baseline and Task-Aware Evidence Modes", 3)
    add_paragraph(
        doc,
        "Two evidence modes are generated for paired evaluation. The baseline_first_20_rows mode "
        "passes the first 20 rows of the completed result together with common task context. This "
        "represents a simple implementation that is easy to reproduce but may omit relevant "
        "evidence when results are longer than 20 rows or poorly ordered for the intended task. "
        "The task_aware_data_summarization mode applies a deterministic strategy selected from the "
        "task metadata. Depending on the task, it may preserve extrema, group summaries, temporal "
        "change, distributions, relationship statistics, risk indicators, or supported action "
        "evidence."
    )
    add_paragraph(
        doc,
        "The task-aware mode is not assumed to be universally superior. Summarization can omit a "
        "detail, emphasize an unhelpful statistic, or create a more complex context than a small "
        "complete result requires. When the full result contains 20 rows or fewer, the baseline "
        "already receives all rows, so differences between modes mainly reflect evidence "
        "organization, prompting, completeness, and claim discipline. When the result exceeds 20 "
        "rows, task-aware summaries have the additional potential advantage of representing "
        "evidence outside the baseline slice. This distinction is retained in the paired analysis."
    )
    add_heading(doc, "3.6.3 Strategy-Based Prompting", 3)
    add_paragraph(
        doc,
        "The explanation strategy is derived from the task rather than chosen freely by the model. "
        "Trend prompts emphasize direction, magnitude, turning points, and temporal coverage. "
        "Comparison prompts emphasize the compared groups and the size of observed differences. "
        "Distribution prompts describe concentration, spread, and notable categories. Correlation "
        "prompts require association language and explicitly prohibit causal wording. Risk prompts "
        "describe observable warning indicators and uncertainty rather than predicting failure. "
        "Behavioral prompts summarize recorded activity without inferring motivation or character. "
        "Ranking prompts contextualize ordered values without equating rank with intrinsic ability."
    )
    add_heading(doc, "3.6.4 Structured Output, Safety, and Traceability", 3)
    add_paragraph(
        doc,
        "The requested response is structured rather than unconstrained prose. The schema includes a "
        "summary, individual insights, evidence references, warnings or limitations, and confidence "
        "information; action-oriented tasks may additionally explain actions already supported by "
        "the evidence. Schema validation rejects malformed responses or routes them to a degraded "
        "fallback. The system records the task, evidence mode, model information, response status, "
        "latency, token usage when available, and error details for later audit."
    )
    add_paragraph(
        doc,
        "Safety constraints prohibit unsupported psychological judgments, deterministic predictions, "
        "and causal claims inferred only from association. Acceptable explanations describe "
        "observable patterns, quote or calculate only supplied numbers, use cautious educational "
        "language, identify incomplete evidence, and avoid identifying individual students in "
        "cohort contexts unless the task is explicitly student-level. For action-synthesis tasks, "
        "the model explains returned actions, triggers, priorities, owners, or time horizons; it is "
        "not required or encouraged to invent additional interventions."
    )
    add_paragraph(
        doc,
        "Traceability follows a deterministic chain: task registry to availability decision, "
        "parameterized SQL execution, normalized query result, evidence packet, model response, "
        "schema validation, and rendered explanation. Evidence identifiers connect statements to "
        "primary or derived evidence. This design does not guarantee that every generated sentence "
        "is correct, but it makes unsupported claims more detectable and provides the artifacts "
        "needed for systematic evaluation."
    )

    add_heading(doc, "3.7 Evaluation Protocol", 2)
    add_paragraph(
        doc,
        "The evaluation protocol tests the enabling infrastructure and the three technical thesis "
        "contributions. Automated runners produce machine-readable outputs, expected-result files "
        "or deterministic checks provide reference evidence, and review artifacts preserve "
        "exceptions and corrections. Supporting evaluations establish that import, mapping, SQL, "
        "API, and visualization behavior are sufficiently stable for the primary evaluations. "
        "Primary evaluations then examine task availability, AI explanation quality, analytical "
        "coverage and human utility, and system performance."
    )
    add_table(
        doc,
        ["Evaluation", "Purpose", "Method and evidence"],
        [
            ["1 Import and conversion", "Verify canonical persistence and integrity", "Row reconciliation, keys, ranges, and engineered-feature checks"],
            ["2 Profiling and mapping", "Verify source detection and mapping behavior", "Profiles, suggestions, previews, and expected mappings"],
            ["3 Task availability", "Test dataset-task validity decisions", "Expected matrix versus actual statuses and diagnostics"],
            ["4 SQL correctness", "Verify analytical calculations", "Task outputs versus deterministic reference queries/checks"],
            ["5 API contract", "Verify stable machine-readable responses", "Schema and error-contract assertions"],
            ["6 Visualization", "Verify metadata-to-chart adaptation", "Semantic-role and rendered-configuration assertions"],
            ["7 AI explanation quality", "Compare two evidence preparation modes", "Paired explanations scored with a fixed rubric and error caps"],
            ["10 Coverage and utility", "Interpret practical analytical reach", "Availability distribution plus reviewer evidence and limitations"],
            ["11 System performance", "Measure operational cost and latency", "Import, validation, SQL, AI, CPU, RAM, token, and cost logs"],
        ],
        [1.5, 2.25, 2.75],
    )
    add_heading(doc, "3.7.1 Datasets and Fixed Task Scope", 3)
    add_paragraph(
        doc,
        "The methodology uses UCI Portuguese student-performance data and OULAD as contrasting "
        "dataset cases. UCI emphasizes compact tabular performance and background variables; OULAD "
        "adds multi-file course, assessment, enrollment, and high-volume temporal engagement data. "
        "The same fixed registry scope is evaluated for both datasets. This design exposes whether "
        "the architecture can distinguish capability differences instead of assuming that a task "
        "supported by one dataset is universally supported."
    )
    add_heading(doc, "3.7.2 Automated Runners and Reproducibility", 3)
    add_paragraph(
        doc,
        "Each evaluation runner records a run identifier, timestamp, dataset scope, configuration or "
        "version information, expected record count, actual outputs, warnings, errors, and summary "
        "statistics. Detailed records are stored in JSON or JSONL, while Markdown reports provide a "
        "human-readable summary. Evaluation outputs are retained under dataset- and phase-specific "
        "directories so that reruns do not silently overwrite earlier evidence. A result is used in "
        "the thesis only after the corresponding runner and validation gate complete successfully "
        "or the report clearly discloses any invalid records."
    )
    add_heading(doc, "3.7.3 Task-Availability Comparison Procedure", 3)
    add_numbered(
        doc,
        [
            "Freeze the task registry and dataset import used for the evaluation.",
            "Construct and review the expected availability record for every dataset-task pair.",
            "Run the validator and capture status, executable flag, warnings, missing requirements, and layer results.",
            "Compare actual and expected outputs and investigate every mismatch rather than relying only on aggregate counts.",
            "Summarize the final status distribution and use it as coverage evidence only after mismatches are resolved or documented.",
        ],
    )
    add_heading(doc, "3.7.4 AI Explanation Scoring Procedure", 3)
    add_paragraph(
        doc,
        "For each executable task in the evaluation scope, the system generates one explanation "
        "using baseline_first_20_rows and one using task_aware_data_summarization. Pairing controls "
        "for the task, dataset, query result, model family, and general output schema. Judge inputs "
        "include task metadata, evidence, full query results when they fit the model context, "
        "deterministic action evidence where applicable, and the generated explanation. When a full "
        "packet exceeds the context limit, a deterministic compact retrieval context is used and "
        "the substitution is logged."
    )
    add_paragraph(
        doc,
        "The scoring rubric covers faithfulness, numerical correctness, completeness, clarity, "
        "actionability, task relevance, and safety/fairness. Metric scores are combined through a "
        "versioned weighted formula. Major errors, such as unsupported numerical claims, can cap "
        "the final score so that fluent wording cannot compensate for a serious grounding failure. "
        "Outputs are validated against the judge-response schema before scoring. Invalid or "
        "internally inconsistent judge responses are retried, repaired only through documented "
        "deterministic procedures, or excluded and reported."
    )
    add_paragraph(
        doc,
        "Paired analysis reports the mean task-aware-minus-baseline score difference, winner and tie "
        "counts, and breakdowns by full-result row-count bucket. The row-count distinction is "
        "methodologically important: for results of at most 20 rows, both modes can receive complete "
        "row coverage; for longer results, the baseline slice can omit evidence. The thesis "
        "therefore interprets average differences together with per-task wins, losses, error types, "
        "and evidence coverage rather than presenting one aggregate as proof of universal "
        "superiority."
    )
    add_heading(doc, "3.7.5 Coverage, Human Utility, and Performance", 3)
    add_paragraph(
        doc,
        "Coverage is derived from the fixed task scope and final availability classifications. "
        "Executable and partial tasks indicate analytical reach, while insufficient and unsupported "
        "tasks reveal the limits imposed by the dataset. Human utility is examined through whether "
        "the system communicates those limits, whether explanations identify relevant evidence and "
        "warnings, and whether a reviewer can trace outputs to their basis. This evaluation is "
        "interpretive and is not presented as a large-scale user study; broader user testing remains "
        "future work."
    )
    add_paragraph(
        doc,
        "Performance evaluation records import duration, availability-validation latency, SQL task "
        "latency, AI request latency, CPU and memory observations where available, and language-model "
        "token and cost information where recoverable. Repeated measurements are summarized using "
        "counts and latency statistics appropriate to each runner. Missing usage metadata is "
        "reported explicitly and is not extrapolated into a complete cost total. Performance results "
        "are interpreted as thesis-environment benchmarks rather than universal production capacity."
    )
    add_heading(doc, "3.7.6 Validity and Ethical Controls", 3)
    add_paragraph(
        doc,
        "Several controls reduce threats to validity. Registry and scoring-formula versions are "
        "frozen within a reported run. Dataset-task ground truth is inspected at the level of "
        "individual requirements. AI modes are paired on the same task outputs. Deterministic "
        "evidence is supplied for action tasks so that the judge evaluates explanation of returned "
        "actions rather than invention of new recommendations. Errors, retries, context compaction, "
        "and missing usage fields are retained in logs. Nevertheless, LLM judge scores remain "
        "model-dependent, and the two datasets do not represent every educational environment."
    )
    add_paragraph(
        doc,
        "Ethically, the platform is limited to descriptive and diagnostic analytics over available "
        "educational records. It does not claim causal inference, psychological profiling, or "
        "deterministic prediction of student success. Availability gating, cohort aggregation, "
        "evidence references, cautious language, and explicit limitations are used to reduce the "
        "risk that missing or correlational evidence is presented as a high-confidence judgment."
    )
    add_heading(doc, "3.8 Chapter Summary", 2)
    add_paragraph(
        doc,
        "This chapter defined the methodology as an evidence chain rather than a collection of "
        "dashboard features. Heterogeneous source files are profiled, human-confirmed, and converted "
        "into a partially populated eight-table canonical schema. Analytical intents are defined in "
        "a metadata-driven task registry. A four-layer validator connects each task to the actual "
        "structural, semantic, analytical, and quantitative capabilities of the imported dataset. "
        "Completed analytical outputs are then transformed into traceable evidence packets for "
        "constrained AI explanation. Finally, automated and paired evaluation protocols measure "
        "validity, explanation quality, coverage, utility, and performance. Chapter 4 describes how "
        "these methodological components are implemented in the software system."
    )


def chapter_4(doc: Document):
    page_break(doc)
    add_heading(doc, "Chapter 4: SYSTEM DESIGN AND IMPLEMENTATION", 1)

    add_heading(doc, "4.1 Overall Architecture", 2)
    add_paragraph(
        doc,
        "The implemented artifact is a three-service learning analytics prototype supported by a "
        "relational database and a file-based analytical task registry. The browser application is "
        "implemented with React and Vite. A Node.js/Express backend coordinates dataset import, "
        "canonical persistence, task discovery, capability validation, SQL execution, and the "
        "proxy connection to the AI service. PostgreSQL stores canonical educational records and "
        "operational provenance through Prisma. A separate Python FastAPI service constructs and "
        "validates evidence-grounded explanations."
    )
    add_paragraph(
        doc,
        "The runtime architecture follows the methodological ordering established in Chapter 3. "
        "Source data are profiled and normalized before analytics are exposed. The active import "
        "batch supplies the dataset context used by the capability validator. Only after a task and "
        "its parameters have been resolved does the SQL execution service run the registry-owned "
        "query. The frontend renders the named result datasets through metadata-driven adapters. "
        "AI explanation is a downstream optional operation over those completed outputs, so an AI "
        "failure cannot invalidate the underlying analytical result."
    )
    add_table(
        doc,
        ["Runtime component", "Technology", "Primary responsibility"],
        [
            ["Web client", "React 19, Vite, React Router, TanStack Query, Ant Design, Recharts", "Role-based workflows, import UI, task selection, deterministic charts, and AI insight display"],
            ["Application API", "Node.js, Express 5, Axios, Multer, csv-parser", "Request validation, orchestration, import, task execution, and AI proxying"],
            ["Persistence", "PostgreSQL with Prisma 7", "Canonical records, import batches, application state, aliases, sessions, and explanation logs"],
            ["AI service", "Python, FastAPI, Pydantic, OpenAI-compatible client", "Strategy selection, evidence summarization, structured response validation, and safety screening"],
            ["Analytical specification", "Version-controlled JSON registry", "Task requirements, SQL, output schema, visualization metadata, and explanation configuration"],
        ],
        [1.3, 2.45, 2.75],
    )
    add_heading(doc, "4.1.1 Component Boundaries", 3)
    add_paragraph(
        doc,
        "The frontend never receives registry SQL and does not construct analytical queries. It "
        "receives a public task view, availability information, named result datasets, and display "
        "metadata. The backend owns source normalization, capability decisions, parameter binding, "
        "and analytical execution. The AI service receives an enriched payload rather than database "
        "credentials. This separation prevents interface components and the language model from "
        "bypassing the system's data and validity controls."
    )
    add_heading(doc, "4.1.2 End-to-End Runtime Sequence", 3)
    add_numbered(
        doc,
        [
            "An administrator uploads one or more CSV files through the import wizard.",
            "The backend profiles the files, detects candidate dataset roles, and returns mapping suggestions.",
            "The user reviews or corrects the mapping, previews transformed rows, and confirms the import.",
            "The backend transforms and persists canonical entities under an ImportBatch, then marks the selected batch active.",
            "The client requests the public task registry and task-availability decisions for the active dataset and class context.",
            "For a selected task, the backend resolves parameters, validates capability, executes registry-owned SQL, and validates the output schema.",
            "The client selects the declared dataset block and chart adapter, applies missing-data policy, and renders the result with diagnostics.",
            "On request, the backend enriches the analytical output with task metadata and sends it to the AI service for structured explanation.",
        ],
    )
    add_paragraph(
        doc,
        "Each transition has a machine-readable contract. Import sessions identify a mapping in "
        "progress; batch identifiers scope persisted data; task identifiers connect the registry to "
        "execution; execution identifiers connect analytics to explanations; query labels connect "
        "multi-query outputs to visualization and evidence roles. This set of identifiers forms the "
        "implementation-level provenance chain."
    )

    add_heading(doc, "4.2 Functional and Non-Functional Requirements", 2)
    add_heading(doc, "4.2.1 Functional Requirements", 3)
    add_table(
        doc,
        ["ID", "Requirement", "Implemented mechanism"],
        [
            ["FR1", "Import heterogeneous educational CSV data", "Upload, profile, mapping review, preview, confirmation, transform, and persistence endpoints"],
            ["FR2", "Maintain dataset context and history", "ImportBatch records, active-dataset state, and batch-scoped canonical rows"],
            ["FR3", "Discover analytical tasks by audience and status", "Public registry view with scope, search, dataset, and governance filters"],
            ["FR4", "Determine task validity before execution", "Four-layer capability validator and diagnostic status response"],
            ["FR5", "Execute approved analytics through one contract", "POST /api/analytics/run with registry-owned SQL and named datasets"],
            ["FR6", "Render results without task-specific chart pages", "Visualization metadata, semantic roles, adapters, and shared ChartRenderer"],
            ["FR7", "Explain completed analytics with AI", "POST /api/ai/explain, strategy factory, evidence summaries, and structured response"],
            ["FR8", "Retain evidence for evaluation", "Import, mapping, availability, explanation, and performance logs"],
        ],
        [0.65, 2.55, 3.3],
    )
    add_heading(doc, "4.2.2 Non-Functional Requirements", 3)
    add_paragraph(
        doc,
        "The most important quality attributes are traceability, data integrity, extensibility, "
        "defensive behavior, and reproducibility. Traceability requires identifiers and logs linking "
        "a result to its batch, task, query configuration, and explanation request. Integrity "
        "requires foreign keys, batch scoping, mapping validation, and output-contract checks. "
        "Extensibility requires that conventional tasks be added through governed metadata rather "
        "than dedicated endpoints. Defensive behavior requires explicit warnings, timeouts, bounded "
        "result sets, null handling, and degraded AI responses. Reproducibility requires that the "
        "registry, evaluation inputs, raw outputs, and reports remain versioned."
    )
    add_table(
        doc,
        ["Quality attribute", "Design response", "Current boundary"],
        [
            ["Performance", "Batch insert, PostgreSQL aggregation, query timeout, result limit, and client-side caching", "Prototype benchmarks; no distributed load test"],
            ["Security", "Server-owned SQL, parameter whitelist, typed binding, CORS, and no direct browser database access", "No full production identity/authorization subsystem"],
            ["Reliability", "Integrity checks, error contracts, output validation, AI degradation, and health endpoints", "External model availability remains a dependency"],
            ["Maintainability", "Route/controller/service separation and metadata-driven task configuration", "Large registry and AI base strategy require future decomposition"],
            ["Privacy and ethics", "Role-aware aggregation, cautious AI rules, evidence constraints, and no causal/predictive claims", "Formal privacy impact assessment is outside the MVP"],
        ],
        [1.25, 3.35, 1.9],
    )

    add_heading(doc, "4.3 Canonical Database Implementation", 2)
    add_paragraph(
        doc,
        "The PostgreSQL schema is defined through Prisma migrations. It implements the eight "
        "canonical learning entities described in Chapter 3: Student, Course, Class, Enrollment, "
        "Assessment, AssessmentResult, Event, and Engagement. Every learning entity includes "
        "batch_id and source_dataset so that identical source identifiers can coexist across "
        "imports without losing provenance. Composite uniqueness rules and indexes combine business "
        "identifiers with batch context where necessary."
    )
    add_heading(doc, "4.3.1 Canonical Relationships", 3)
    add_paragraph(
        doc,
        "Course is connected to one or more class presentations. Student and Class are connected "
        "through Enrollment. Assessment belongs to the instructional context, while "
        "AssessmentResult connects a learner or enrollment to an assessment outcome. Event defines "
        "a learning resource or activity, and Engagement records interactions within an enrollment "
        "and batch context. These relationships support performance, outcome, demographic, and "
        "activity queries while preserving referential integrity."
    )
    add_table(
        doc,
        ["Entity group", "Models", "Key implementation role"],
        [
            ["Learning context", "Course, Class", "Represents modules and specific presentations/runs"],
            ["Participation", "Student, Enrollment", "Connects learners to classes and stores registration/final outcome context"],
            ["Performance", "Assessment, AssessmentResult", "Stores graded activities, score evidence, and submission information"],
            ["Activity", "Event, Engagement", "Stores resource metadata and batch-scoped temporal interaction counts"],
        ],
        [1.35, 2.15, 3.0],
    )
    add_heading(doc, "4.3.2 Operational and Provenance Tables", 3)
    add_paragraph(
        doc,
        "The implementation adds operational models that are separate from the educational domain. "
        "ImportBatch records dataset identity, source type, sample status, import status, row counts, "
        "and whether the batch is active. AppState stores the current application-level selection. "
        "UploadSession retains the intermediate profiling and mapping state. AliasMemory stores "
        "confirmed source-to-canonical aliases used by future suggestions. AiExplanationLog records "
        "task, strategy, execution linkage, response status, latency, and degraded-state information. "
        "This separation prevents workflow metadata from being confused with learning evidence."
    )
    add_heading(doc, "4.3.3 Feature Persistence and Runtime Metrics", 3)
    add_paragraph(
        doc,
        "Persisted engineered fields follow the within-table rule. Student stores contextual scores "
        "such as lifestyle_risk_score, support_score, social_balance_score, "
        "family_stability_score, and disadvantage_score when their source inputs are available. "
        "Enrollment stores registration_lead_time; Assessment stores week_of_class; "
        "AssessmentResult stores pass_flag; and Engagement stores log_click_score. Cross-table "
        "quantities such as class average, pass rate, total clicks, score trend, correlation, and "
        "composite task-specific risk indicators remain SQL runtime outputs."
    )
    add_heading(doc, "4.3.4 Batch Isolation and Deletion", 3)
    add_paragraph(
        doc,
        "Batch isolation is essential because the system can load UCI, OULAD, sample data, and "
        "future custom datasets into one database. Queries and capability checks include batch "
        "context, and many canonical relationships are indexed by batch. ImportBatch relationships "
        "use cascading deletion where appropriate so that removing one imported dataset does not "
        "leave orphan analytical records. The active flag selects the default working dataset but "
        "does not remove historical batches."
    )

    add_heading(doc, "4.4 Import Pipeline Implementation", 2)
    add_paragraph(
        doc,
        "The import subsystem is exposed through four main endpoints: POST /api/import/profile, "
        "POST /api/import/confirm-mapping, POST /api/import/preview, and POST /api/import/run. "
        "Separating these operations allows the user to inspect the system's interpretation before "
        "records are committed. The frontend implements this as upload, review, confirmation, and "
        "completion steps rather than a single opaque file-upload action."
    )
    add_heading(doc, "4.4.1 Profiling and Source Detection", 3)
    add_paragraph(
        doc,
        "The profile controller receives CSV files through Multer, detects the delimiter, parses raw "
        "rows, and computes a profile containing headers, inferred types, example values, missing "
        "values, and row information. Dataset-type and file-role detectors compare filename and "
        "column signatures with known UCI and OULAD rules. For multi-file uploads, bundle-schema "
        "detection identifies the likely set of related educational entities. The resulting profile, "
        "raw rows, suggestions, and session metadata are stored in an UploadSession after temporary "
        "upload files are removed."
    )
    add_heading(doc, "4.4.2 Mapping Suggestion, Confirmation, and Preview", 3)
    add_paragraph(
        doc,
        "Mapping suggestions combine source-name normalization, canonical aliases, type "
        "compatibility, sample-value patterns, known dataset rules, and AliasMemory. The review UI "
        "allows a user to change the target entity, target field, or transformation. Confirmation "
        "sets the mapping state to confirmed and invokes strict mapping validation. Manual "
        "corrections can be persisted as learned aliases when they represent a previously unknown "
        "source label. The preview endpoint runs transformation on a small sample and returns "
        "canonical rows without committing them."
    )
    add_heading(doc, "4.4.3 Transformation and Identifier Construction", 3)
    add_paragraph(
        doc,
        "The runImportPipeline service invokes transformRawRowsToCanonical. Flat UCI records are "
        "decomposed into students, courses, classes, enrollments, assessments, and assessment "
        "results. Multiple grade columns are unpivoted into assessment records. Multi-file OULAD "
        "records are integrated through module, presentation, student, assessment, and resource "
        "keys. Surrogate identifiers include enough batch and instructional context to avoid "
        "collisions. Vocabulary and score transformations are applied through canonical rules, and "
        "within-table features are calculated only when their inputs are present."
    )
    add_heading(doc, "4.4.4 Persistence and Large-File Handling", 3)
    add_paragraph(
        doc,
        "Canonical entity arrays are deduplicated before insertion. The service inserts Course, "
        "Class, Student, Enrollment, Assessment, AssessmentResult, Event, and Engagement in "
        "foreign-key-safe order. Prisma createMany with skipDuplicates is used for conventional "
        "batches, while PostgreSQL COPY support is available for high-volume paths such as OULAD "
        "engagement. On successful completion, an ImportBatch is recorded and can be activated. "
        "The use of skipDuplicates prevents insertion failure but makes duplicate reporting "
        "important; this is retained as an implementation limitation rather than interpreted as "
        "proof that the source contained no duplicates."
    )
    add_heading(doc, "4.4.5 Import Diagnostics", 3)
    add_paragraph(
        doc,
        "The pipeline emits raw and canonical row counts, duplicate/skipped information where "
        "available, foreign-key and range checks, and engineered-feature checks. Profiling and "
        "mapping logs preserve detection and suggestion behavior, while import-conversion logs "
        "preserve persistence outcomes. These artifacts are used by the supporting evaluations in "
        "Chapter 5 and make transformation failures distinguishable from later analytics failures."
    )

    add_heading(doc, "4.5 Task Registry and Availability Validator", 2)
    add_paragraph(
        doc,
        "The registry is stored in Backend/src/config/taskRegistry.json and loaded through "
        "TaskRegistryService. At the time of this draft, the implementation registry contains 57 "
        "task records. The thesis evaluations use a frozen 52-task scope after applying evaluation "
        "eligibility and run-specific controls; the distinction prevents later experimental or "
        "non-evaluated records from silently changing reported denominators."
    )
    add_heading(doc, "4.5.1 Registry Loading and Public Views", 3)
    add_paragraph(
        doc,
        "The service loads registry JSON once, strips SQL comments for validation, detects malformed "
        "multi-statement definitions, and supports filtering by scope, dataset, keyword, and analysis "
        "type. Governance status controls exposure: deprecated tasks are omitted; experimental tasks "
        "require explicit inclusion; validated and certified tasks are exposed by default. Before "
        "metadata are sent to the browser, the controller removes sqlQuery and sqlQueries so that "
        "server implementation details are not leaked into the client contract."
    )
    add_table(
        doc,
        ["Metadata group", "Representative fields", "Runtime consumer"],
        [
            ["Identity and governance", "taskId, taskName, scope, roles, registry_status", "Registry service and task browser"],
            ["Data requirements", "sourceTables, keyDbFields, requiredCapabilities, availability_contract", "Capability validator"],
            ["Execution", "sqlQuery/sqlQueries, params, query_labels, output_schema", "SQL executor and output validator"],
            ["Visualization", "viz_type, visualization_config, semantic_roles", "ChartRenderer and adapters"],
            ["Explanation", "analysis_context, explanation_strategy, target_audience, AI summary configuration", "Node AI proxy and FastAPI strategy"],
        ],
        [1.45, 3.05, 2.0],
    )
    add_heading(doc, "4.5.2 Capability Snapshot", 3)
    add_paragraph(
        doc,
        "CanonicalCapabilityService builds a batch-aware snapshot by querying populated canonical "
        "fields and counts. The snapshot can identify assessment scores, final outcomes, submission "
        "history and timestamps, demographics, lifestyle and family context, socioeconomic context, "
        "engagement tracking, resource clickstream, temporal activity, absences, study time, "
        "registration timing, and multi-student comparison support. Native competency tags and "
        "proxy competency availability are distinguished to avoid presenting correlated grades as a "
        "formal competency ontology."
    )
    add_heading(doc, "4.5.3 Layer A-D Validation", 3)
    add_paragraph(
        doc,
        "Layer A queries PostgreSQL information_schema and checks that registry source tables are in "
        "an allowed canonical set. Layer B evaluates the task's required_all, required_any, "
        "optional_enrichments, dataset-specific, and engineered-feature conditions against the "
        "capability snapshot. Layer C checks analytical suitability, such as at least two temporal "
        "points for a trend or at least two students for comparison; current Layer C issues are "
        "advisory. Layer D queries enrollment, result, diversity, week, and engagement counts against "
        "minimum conditions."
    )
    add_paragraph(
        doc,
        "The validator returns layer results, status, executable flag, confidence, warnings, and "
        "missing requirements. Structural failure maps to unsupported; data-sufficiency failure maps "
        "to insufficient_data; semantic failure maps to partial; otherwise the task is executable. "
        "Confidence is HIGH for stronger student, assessment, and temporal coverage, MEDIUM for the "
        "minimum ordinary cohort conditions, and LOW when the task can run with limited evidence."
    )
    add_heading(doc, "4.5.4 Execution Policy and UI Communication", 3)
    add_paragraph(
        doc,
        "The current analytics endpoint hard-blocks unsupported tasks with an HTTP 422 response. "
        "Partial and insufficient_data statuses are soft failures: execution may continue with data "
        "quality warnings so that valid degraded outputs and diagnostic evaluation remain possible. "
        "The response embeds status, confidence, confidence reason, and warnings. The frontend uses "
        "availability messages and badges to expose these limitations. This is a deliberate MVP "
        "trade-off; a production policy could additionally hard-block selected insufficient-data "
        "tasks through per-task execution contracts."
    )

    add_heading(doc, "4.6 SQL Analytics Execution", 2)
    add_paragraph(
        doc,
        "All analytical SQL originates from the server-side task registry. The request supplies a "
        "task identifier and permitted values, not arbitrary SQL text. POST /api/analytics/run is "
        "therefore a shared execution endpoint whose behavior is selected by governed metadata. "
        "The controller creates an executionId, resolves student/class/enrollment context, loads the "
        "task and batch, runs capability validation, and invokes SqlExecutionService."
    )
    add_heading(doc, "4.6.1 Parameter Resolution and Guardrails", 3)
    add_paragraph(
        doc,
        "When a client provides student_id and class_id but not enrollment_id, the controller "
        "resolves the enrollment within the batch context. SQL placeholders are restricted to a "
        "whitelist that includes batch, class, student, enrollment, and two-student comparison "
        "parameters. Placeholders are converted to PostgreSQL positional parameters and bound rather "
        "than interpolated into the SQL string. Unknown request parameters or unknown placeholders "
        "are rejected. Registry SQL comments are stripped and multi-statement misuse is detected at "
        "load time."
    )
    add_table(
        doc,
        ["Guardrail", "Implementation", "Purpose"],
        [
            ["Query ownership", "SQL loaded only from task registry", "Prevents user- or model-authored arbitrary SQL"],
            ["Parameter control", "Whitelist plus positional binding", "Reduces SQL injection and context errors"],
            ["Batch scoping", "Batch-aware filters and join rewrites", "Prevents records from different imports being mixed"],
            ["Timeout", "SET LOCAL statement_timeout = 30 seconds inside transaction", "Bounds long-running query impact"],
            ["Result bound", "LIMIT 10000 injected when absent", "Bounds response and memory volume"],
            ["Output contract", "Required columns validated for non-empty datasets", "Detects registry-query/frontend drift"],
        ],
        [1.25, 3.0, 2.25],
    )
    add_heading(doc, "4.6.2 Single- and Multi-Query Tasks", 3)
    add_paragraph(
        doc,
        "A registry record may contain one sqlQuery or an ordered sqlQueries collection. Single-query "
        "results are assigned to datasets[query_label]. Multi-query results are assigned to separate "
        "named blocks using the corresponding query_labels. Named blocks are preferable to a flat "
        "data array because a task may require, for example, a summary dataset and a trend dataset "
        "with different schemas. The same labels are later used by chart selection and AI evidence "
        "roles."
    )
    add_heading(doc, "4.6.3 Result Normalization and Diagnostics", 3)
    add_paragraph(
        doc,
        "Database BigInt and Decimal values are converted to JSON-safe values. The service records "
        "execution time and computes a queryHash from the task and parameters for traceability. The "
        "output-schema validator checks required columns when a dataset is non-empty; a mismatch "
        "returns OUTPUT_SCHEMA_MISMATCH with HTTP 422 rather than allowing the frontend to guess. "
        "The successful response includes executionId, taskId, named datasets, row counts, query "
        "labels, execution timing, query hash, and data-quality metadata."
    )

    add_heading(doc, "4.7 Frontend and Visualization Layer", 2)
    add_paragraph(
        doc,
        "The frontend is organized around role selection, dataset import, dataset selection, student "
        "and administrator dashboards, and a general analytics workspace. React Router lazy-loads "
        "these pages. AppContext stores active-dataset and import-history state, while TanStack Query "
        "manages server state with a five-minute global stale time, one retry, and no automatic "
        "refetch on window focus. This reduces unnecessary reruns during interactive analysis."
    )
    add_heading(doc, "4.7.1 Role-Based Workspaces", 3)
    add_paragraph(
        doc,
        "Administrators can import and activate datasets, review cohort-oriented tasks, and use the "
        "general analytics workspace. Students use a personal dashboard that resolves active dataset, "
        "class, enrollment, and compatible student tasks. Both roles share task metadata, analytics "
        "hooks, ChartRenderer, raw-result fallback, and AIInsightPanel. The difference lies primarily "
        "in task scope, parameter selection, default task sets, and aggregation level rather than in "
        "separate analytical engines."
    )
    add_heading(doc, "4.7.2 Metadata-Driven Chart Rendering", 3)
    add_paragraph(
        doc,
        "ChartRenderer reads taskMeta.viz_type and visualization_config, selects the relevant named "
        "dataset, resolves required semantic fields, invokes an adapter, and then selects a visual "
        "component. Supported types include line chart, bar chart, histogram, scatter plot, pie "
        "chart, heatmap, table, metric card, and checklist. The adapter boundary decouples API row "
        "shape from the props expected by Recharts or custom components."
    )
    add_table(
        doc,
        ["Visualization", "Typical use", "Adapter responsibility"],
        [
            ["Line chart", "Assessment or engagement trends", "Map temporal/category and numeric series fields"],
            ["Bar/histogram", "Group comparison or distribution", "Group rows, preserve categories, and validate numeric values"],
            ["Scatter plot", "Relationship analysis", "Extract paired finite x/y values and labels"],
            ["Heatmap", "Activity intensity over two dimensions", "Build valid coordinate cells and value scale"],
            ["Card/checklist", "Summary metrics, flags, or returned actions", "Map semantic metrics and active/triggered states"],
            ["Table", "Detailed or heterogeneous analytical rows", "Preserve columns and format values for inspection"],
        ],
        [1.35, 2.5, 2.65],
    )
    add_heading(doc, "4.7.3 Missing-Data and Diagnostic Policy", 3)
    add_paragraph(
        doc,
        "The shared adapter policy preserves real zero and does not silently convert null, missing, "
        "or non-finite values to zero. Invalid rows are skipped with counts. Diagnostics record input "
        "rows, valid rows, skipped rows, missing fields, per-field missing counts, selected dataset "
        "label, null-handling policy, and warnings. If no valid data remain, the renderer shows an "
        "explicit empty-state message. Diagnostics can be exposed in the interface during review and "
        "evaluation, which makes visualization loss observable rather than hidden."
    )
    add_heading(doc, "4.7.4 Task Availability in the Interface", 3)
    add_paragraph(
        doc,
        "Task lists and detail panels display availability status and reasons. Unsupported tasks are "
        "not presented as ordinary runnable analyses. Partial and insufficient-data tasks can carry "
        "warnings into result presentation. A raw table or JSON result remains available as a "
        "fallback when a specialized chart cannot be constructed, allowing execution correctness to "
        "be inspected independently of chart configuration. Some frontend compatibility hints still "
        "exist for usability and are identified as a source of possible drift from the backend "
        "validator in the limitations section."
    )

    add_heading(doc, "4.8 AI Explanation Service", 2)
    add_paragraph(
        doc,
        "AI explanation is implemented as a Node-to-Python request chain. The browser calls POST "
        "/api/ai/explain with the task identifier, analytics execution identifier, and named result "
        "datasets. The Node controller reloads the server-side task definition and enriches the "
        "payload with explanation strategy, target audience, analysis context, visualization "
        "configuration, query labels, evidence columns, dataset roles, and action-evidence contracts. "
        "It then calls the FastAPI /explain endpoint with a configured timeout."
    )
    add_heading(doc, "4.8.1 Strategy Factory and Evidence Summarization", 3)
    add_paragraph(
        doc,
        "FastAPI validates the request through Pydantic and resolves the declared strategy through "
        "ExplanationStrategyFactory. Registered strategies include trend, comparison, distribution, "
        "correlation, risk, behavioral, ranking, recommendation, and progress-oriented variants. A "
        "base strategy implements shared evidence handling and task-aware summary configurations. "
        "Specialized strategy prompts constrain the interpretation required for the analytical type. "
        "An unknown strategy raises an error that is converted into a degraded response rather than "
        "falling back to unconstrained generic prompting."
    )
    add_heading(doc, "4.8.2 Action and Derived Evidence", 3)
    add_paragraph(
        doc,
        "Action-oriented tasks require special handling because the analytical query or deterministic "
        "backend rule may already return a recommended action, flag, priority, owner, or time horizon. "
        "The AI summary configuration identifies evidence columns, primary and derived evidence "
        "roles, trigger conditions, sensitive fields, and missing-evidence behavior. The model is "
        "asked to explain the returned action logic rather than create a new intervention. Derived "
        "evidence records retain links to their source evidence identifiers."
    )
    add_heading(doc, "4.8.3 Structured Response and Safety Screening", 3)
    add_paragraph(
        doc,
        "The AI service requests structured JSON and builds an ExplainResponse containing summary, "
        "insights, evidence, recommendations when contractually appropriate, warnings, confidence, "
        "strategy metadata, and safety flags. Pydantic enforces the response shape. A rule-based "
        "SafetyFilter scans generated string fields for prohibited categories such as deterministic "
        "prediction, personality judgment, unsupported psychological inference, or causal language. "
        "Flags are retained in the response and logs for review."
    )
    add_heading(doc, "4.8.4 Graceful Degradation and Logging", 3)
    add_paragraph(
        doc,
        "If the Python service, external model, strategy, validation, or timeout path fails, either "
        "FastAPI or the Node proxy builds a degraded response with degraded=true, a reason, warnings, "
        "and no fabricated explanation. The Node proxy returns HTTP 200 for this valid degraded UI "
        "state so the analytical result remains visible and the client does not treat optional AI "
        "unavailability as total request failure. AiExplanationLog records the linked task and "
        "execution, strategy, latency, degraded status, and failure reason. Token and cost metadata "
        "are exposed when the provider returns them."
    )

    add_heading(doc, "4.9 API and Output Contracts", 2)
    add_paragraph(
        doc,
        "The Express API is divided into import, datasets, tasks, analytics, AI, students, and classes "
        "route groups, with /api/health for runtime checks. Swagger/OpenAPI configuration defines "
        "reusable request and response schemas, while route-level documentation describes the "
        "endpoints. The principal contracts are TaskMetadata, TaskValidationResult, "
        "AnalyticsRequest/AnalyticsResponse, VisualizationConfig and SemanticRoles, and the AI "
        "request, explanation, evidence, confidence, metadata, and degraded-response schemas."
    )
    add_table(
        doc,
        ["Endpoint group", "Representative route", "Contract role"],
        [
            ["Import", "POST /api/import/profile|preview|confirm-mapping|run", "Separates inspection, confirmation, and persistence"],
            ["Dataset context", "GET /api/datasets; classes; students", "Provides active batch and parameter selections"],
            ["Tasks", "GET /api/tasks; /validate/:datasetId; /validate-one/:taskId", "Provides public metadata and validity diagnostics"],
            ["Analytics", "POST /api/analytics/run", "Returns named datasets, execution metadata, and data quality"],
            ["AI", "POST /api/ai/explain", "Returns structured explanation or valid degraded response"],
            ["System", "GET /api/health; /api-docs", "Supports service probes and interactive contract inspection"],
        ],
        [1.25, 3.0, 2.25],
    )
    add_heading(doc, "4.9.1 Analytics Response Contract", 3)
    add_paragraph(
        doc,
        "A successful analytics response contains success, executionId, taskId, datasets, and meta. "
        "datasets is an object keyed by query_labels rather than one untyped array. meta includes "
        "whether the task is multi-query, row counts, execution time, query hash, labels, and a "
        "dataQuality object with status, confidence, confidence reason, and warnings. This shape is "
        "consumed consistently by the analytics workspace, dashboards, visualization adapter, AI "
        "payload builder, and evaluation runners."
    )
    add_heading(doc, "4.9.2 Error and Degraded Contracts", 3)
    add_paragraph(
        doc,
        "Technical and contract failures use a standardized error body with success=false, an error "
        "code, message, and contextual details where safe. Examples include missing task parameters, "
        "unsupported capability, query failure, and output-schema mismatch. AI unavailability is "
        "modeled differently because it is an optional downstream feature: a structured degraded "
        "response is returned successfully and rendered as a limitation. This distinction allows "
        "clients to separate analytical failure from explanation unavailability."
    )

    add_heading(doc, "4.10 Evaluation and Logging Tooling", 2)
    add_paragraph(
        doc,
        "Evaluation tooling is implemented alongside the runtime rather than reconstructed from "
        "screenshots. Backend services write import-conversion, profiling-mapping, and task-availability "
        "records. Dedicated scripts validate registry structure, calculate registry statistics, run "
        "tasks against datasets, freeze registry state, and benchmark performance. Frontend tests and "
        "verification scripts exercise chart adapters, dashboard task sets, URL state, risk cards, "
        "checklists, and experimental-task visibility."
    )
    add_heading(doc, "4.10.1 Evaluation Artifact Structure", 3)
    add_paragraph(
        doc,
        "Primary evaluation runs use phase- and dataset-specific directories. AI evaluation preserves "
        "generation inputs, baseline and task-aware explanations, judge inputs, compact or full judge "
        "contexts, raw judge outputs, validation records, scoring manifests, final scoring records, "
        "aggregate comparisons, and human-readable reports. JSONL is used for record-level appendable "
        "artifacts, while JSON and Markdown reports summarize validated outputs. This structure makes "
        "it possible to inspect one task pair without rerunning all 208 explanation records."
    )
    add_heading(doc, "4.10.2 Performance Instrumentation", 3)
    add_paragraph(
        doc,
        "The systemPerformanceBenchmark script measures task-availability and SQL execution paths for "
        "the available datasets and records latency, failure rate, CPU, and memory observations where "
        "supported. Import logs provide separate conversion duration and row-volume evidence. AI and "
        "judge invocation logs capture provider latency and usage metadata when recoverable. Query "
        "hashes, execution identifiers, run identifiers, timestamps, dataset scopes, and model names "
        "allow performance observations to be associated with the correct configuration."
    )
    add_heading(doc, "4.10.3 Debug Agent", 3)
    add_paragraph(
        doc,
        "A one-click debug agent under agents/ checks environment, database connectivity, registry "
        "loading, validator behavior, analytics execution, output contracts, frontend verification, "
        "and evaluation artifacts. It is intentionally report-only and is not part of the ordinary "
        "user flow. The agent does not edit source code automatically, which keeps diagnostic "
        "evidence separate from remediation and prevents evaluation runs from mutating the system "
        "they are intended to inspect."
    )

    add_heading(doc, "4.11 Implementation Trade-offs and Limitations", 2)
    add_paragraph(
        doc,
        "The implementation prioritizes a complete, auditable thesis-scale pipeline over distributed "
        "production optimization. Queries execute on demand against the canonical database rather "
        "than a warehouse or precomputed analytical mart. Registry SQL placeholder handling is "
        "controlled by whitelist and binding but still uses parsing rules appropriate to server-owned "
        "templates, not arbitrary user SQL. UploadSession can retain raw rows in JSON during the "
        "review workflow, which is convenient for moderate files but unsuitable for unlimited "
        "uploads. BigInt conversion to ordinary JSON numbers could lose precision at extreme values."
    )
    add_paragraph(
        doc,
        "The task registry is centralized but increasingly large, and some frontend compatibility "
        "hints duplicate backend availability knowledge. Table visualization accepts heterogeneous "
        "shapes more permissively than specialized charts. Severe chart row loss is logged in "
        "diagnostics but could be promoted more prominently in a production interface. The AI service "
        "has degraded fallbacks, yet non-degraded quality and cost still depend on an external model. "
        "These limitations motivate registry validation in continuous integration, storage-backed "
        "upload streaming, stricter unit metadata, a single backend capability source of truth, and "
        "larger-scale user and load testing."
    )

    add_heading(doc, "4.12 Chapter Summary", 2)
    add_paragraph(
        doc,
        "This chapter described the concrete implementation of the methodology. A React client and "
        "Express API coordinate a batch-scoped PostgreSQL canonical model, a human-confirmed import "
        "pipeline, a governed analytical task registry, a four-layer capability validator, bounded "
        "parameterized SQL execution, metadata-driven chart adapters, and a separate FastAPI "
        "explanation service. Stable identifiers and contracts connect import evidence, capability "
        "decisions, analytical outputs, visualizations, and AI explanations. Logging and evaluation "
        "tools preserve this chain for empirical analysis. Chapter 5 uses those artifacts to report "
        "the system's verification, task availability, explanation quality, coverage, utility, and "
        "performance results."
    )


def chapter_5_outline(doc: Document, evidence: dict):
    page_break(doc)
    add_heading(doc, "Chapter 5: EXPERIMENTS AND RESULTS", 1)
    add_heading(doc, "5.1 Experimental Setup", 2)
    add_paragraph(
        doc,
        "The experiments use UCI Student Performance and OULAD as representative heterogeneous "
        "educational datasets. UCI provides student demographic, family, school, lifestyle, and "
        "assessment-related attributes. OULAD provides course, assessment, enrollment, and large "
        "virtual learning environment interaction data. These datasets are intentionally different, "
        "making them suitable for evaluating the reuse and validity claims of the proposed architecture."
    )
    add_paragraph(
        doc,
        "The evaluation is organized around eleven internal evaluation areas. This chapter focuses "
        "on the primary evidence areas recommended for thesis framing: Evaluation 3, Evaluation 7, "
        "Evaluation 10, and Evaluation 11. Other evaluations are summarized as supporting checks."
    )

    add_heading(doc, "5.2 Supporting System Verification", 2)
    ic = evidence["import_conversion"]
    add_paragraph(
        doc,
        "The supporting verification checks confirm that the system is stable enough for the main "
        "evaluation. The UCI import conversion log reports a successful import with zero foreign-key "
        "errors and zero null/range errors. Feature checks also report zero errors for pass flags, "
        "week-of-class derivation, registration lead time, log-click score, and student composite features."
    )
    rows = []
    for k, v in ic["canonical_rows"].items():
        rows.append([k, str(v)])
    add_table(doc, ["Canonical table", "Rows inserted in UCI import"], rows, [3.0, 3.0])
    add_source_note(doc, "Source: Docs/evaluation_logs/import_conversion/import_conversion_20260620T080702Z_043fc1.json.")

    add_heading(doc, "5.3 Evaluation 3: Task Availability", 2)
    add_paragraph(
        doc,
        "Task availability evaluation tests whether the validator correctly classifies analytical "
        "tasks according to the evidence available in each dataset. This is central to the thesis "
        "because a reusable dashboard should not expose all tasks indiscriminately after import."
    )
    ta = evidence["task_availability"]
    add_table(
        doc,
        ["Dataset", "Total tasks", "Executable", "Partial", "Insufficient data", "Unsupported"],
        [
            [
                "UCI Portuguese",
                str(ta["UCI Portuguese"]["total"]),
                str(ta["UCI Portuguese"]["executable"]),
                str(ta["UCI Portuguese"]["partial"]),
                str(ta["UCI Portuguese"]["insufficient_data"]),
                str(ta["UCI Portuguese"]["unsupported"]),
            ],
            [
                "OULAD",
                str(ta["OULAD"]["total"]),
                str(ta["OULAD"]["executable"]),
                str(ta["OULAD"]["partial"]),
                str(ta["OULAD"]["insufficient_data"]),
                str(ta["OULAD"]["unsupported"]),
            ],
        ],
        [1.5, 1.0, 1.0, 1.0, 1.2, 1.0],
    )
    add_source_note(doc, "Sources: task_availability_20260619T042411Z_265c90.json and task_availability_20260619T042455Z_0ef6b3.json.")
    add_paragraph(
        doc,
        "The results show that OULAD supports more executable tasks than UCI Portuguese in the "
        "fixed scope. This is expected because OULAD contains richer temporal engagement and "
        "assessment interaction data. UCI still supports a meaningful subset of analyses, but many "
        "tasks become partial or insufficient because required temporal or engagement capabilities "
        "are missing. This confirms the need for task availability validation rather than a single "
        "static dashboard menu."
    )

    add_heading(doc, "5.4 Evaluation 7: AI Explanation Quality", 2)
    ai = evidence["ai_scoring"]
    add_paragraph(
        doc,
        "AI explanation quality is evaluated by comparing a baseline first-20-rows context strategy "
        "with a task-aware data summarization strategy. The judge scores each explanation using "
        "metrics for faithfulness, numerical correctness, completeness, clarity, actionability, "
        "task relevance, and safety/fairness. The goal is not to prove that task-aware summarization "
        "always wins, but to determine whether it provides a more reliable evidence organization "
        "strategy for dashboard explanations."
    )
    add_table(
        doc,
        ["Dataset", "Scored records", "Average final score", "Average delta", "Winner counts"],
        [
            [
                "UCI Portuguese",
                ai["UCI Portuguese"]["records"],
                ai["UCI Portuguese"]["avg_final"],
                ai["UCI Portuguese"]["avg_delta"],
                ai["UCI Portuguese"]["winners"],
            ],
            [
                "OULAD",
                ai["OULAD"]["records"],
                ai["OULAD"]["avg_final"],
                ai["OULAD"]["avg_delta"],
                ai["OULAD"]["winners"],
            ],
        ],
        [1.3, 1.0, 1.2, 1.0, 2.0],
    )
    add_source_note(doc, "Sources: official scoring reports for SAMPLE_UCI_POR and SAMPLE_OULAD in Docs/evaluation_v2/Runs/full_208/.")
    add_paragraph(
        doc,
        "The average delta is approximately +0.09 for both datasets. This means task-aware "
        "summarization improves the baseline slightly on average, but the improvement is not "
        "large enough to claim that it is universally superior. The paired comparison shows ties "
        "and baseline wins as well as task-aware wins. The strongest task-aware advantages occur "
        "when the full SQL result contains more than 20 rows and the baseline risks unsupported "
        "claims from an incomplete evidence slice."
    )
    add_paragraph(
        doc,
        "This result supports a careful interpretation: the contribution is not that AI always "
        "becomes more accurate, but that the system provides an auditable evidence protocol for "
        "AI explanation. The protocol reveals both improvements and failure cases."
    )

    add_heading(doc, "5.5 Evaluation 10: Coverage and Human Utility", 2)
    add_paragraph(
        doc,
        "Coverage and human utility are evaluated by examining how many tasks are available, "
        "which task categories are affected by missing data, and whether the dashboard communicates "
        "the result in a way that is useful for human interpretation. The task availability results "
        "serve as quantitative coverage evidence, while the review of AI explanations and disabled-task "
        "messages provides qualitative utility evidence."
    )
    add_paragraph(
        doc,
        "The UCI result demonstrates the value of partial and insufficient-data states. A dashboard "
        "without availability validation might simply fail at runtime or hide why an engagement task "
        "cannot be executed. The proposed system instead exposes the reason. The OULAD result shows "
        "that the same taxonomy can unlock more tasks when richer engagement evidence exists. Thus, "
        "the taxonomy is reusable, but the executable subset remains dataset-specific."
    )
    add_paragraph(
        doc,
        "Human utility is strongest when the system turns technical missing-data conditions into "
        "actionable explanations. For example, an administrator can see that a task is unavailable "
        "because engagement tracking is absent, rather than interpreting the absence of a chart as "
        "a system error. This improves transparency even when the system cannot execute the desired "
        "analysis."
    )

    add_heading(doc, "5.6 Evaluation 11: System Performance", 2)
    add_paragraph(
        doc,
        "System performance evaluation measures local runtime behavior for task availability, SQL "
        "analytics, AI explanation, import, and resource usage. The benchmark uses sequential measured "
        "requests, warm-up runs, CPU/RAM sampling, and JSON performance logs."
    )
    add_table(
        doc,
        ["Scenario", "Dataset", "Average", "P95", "Error rate"],
        [
            ["Task Availability", "UCI", "841.282 ms", "1,249.040 ms", "0%"],
            ["Task Availability", "OULAD", "28,790.822 ms", "31,048.256 ms", "0%"],
            ["Simple Analytics A-B02", "UCI", "16.761 ms", "26.149 ms", "0%"],
            ["Simple Analytics A-B02", "OULAD", "66.890 ms", "102.880 ms", "0%"],
            ["Trend Analytics S-T01", "UCI", "59.235 ms", "71.140 ms", "0%"],
            ["Trend Analytics S-T01", "OULAD", "107.205 ms", "138.737 ms", "0%"],
            ["AI Explanation", "UCI", "5,733.360 ms", "6,284.419 ms", "0%"],
            ["AI Explanation", "OULAD", "6,075.344 ms", "7,420.751 ms", "0%"],
        ],
        [1.8, 1.0, 1.2, 1.2, 1.0],
    )
    add_source_note(doc, "Source: Docs/evaluation_logs/system_performance/SYSTEM_PERFORMANCE_EVALUATION.md.")
    add_paragraph(
        doc,
        "SQL analytics remain fast on both datasets. The main backend bottleneck is OULAD task "
        "availability, which must evaluate the availability of many tasks against a richer and larger "
        "dataset. AI explanation is slower than SQL analytics because it includes an external model "
        "call. The measured AI cost for 20 explanation calls is approximately USD 0.010241."
    )
    add_table(
        doc,
        ["Import scenario", "Rows / scope", "Duration", "Rows/sec", "Error rate"],
        [
            ["UCI import pipeline", "6,490 input rows across 10 measured imports", "2,408.664 ms average", "289.229", "0%"],
            ["OULAD full import", "10,655,280 raw studentVle rows; 8,459,320 engagement rows inserted", "1,028,042.561 ms (17.13 min)", "10,364.629", "0%"],
        ],
        [1.4, 2.4, 1.4, 0.8, 0.8],
    )
    add_paragraph(
        doc,
        "The OULAD full import result is important because it demonstrates that the canonical "
        "pipeline can process a large clickstream file at thesis scale. The benchmark is not a "
        "production load test, but it supports the claim that the architecture is practically usable "
        "on the local evaluation machine."
    )

    add_heading(doc, "5.7 Summary of Experimental Findings", 2)
    add_bullets(
        doc,
        [
            "RQ1 is supported by successful import and canonical conversion logs, especially the UCI import with zero integrity and feature-check errors.",
            "RQ2 is supported by task availability results that differ meaningfully between UCI and OULAD while using the same fixed task scope.",
            "RQ3 is partially supported: task-aware AI explanations score highly and improve the baseline slightly on average, but they do not always win.",
            "RQ4 is supported by coverage and performance logs showing useful task coverage, transparent unavailable-task handling, stable analytics execution, and acceptable local performance.",
        ],
    )


def chapter_5(doc: Document, evidence: dict):
    page_break(doc)
    add_heading(doc, "Chapter 5: EXPERIMENTS AND RESULTS", 1)

    add_heading(doc, "5.1 Experimental Setup", 2)
    add_paragraph(
        doc,
        "This chapter reports the evidence used to assess the three technical contributions: the "
        "canonical schema, the task taxonomy and availability mechanism, and the evidence-grounded "
        "AI explanation pipeline. Evaluation is not treated as a fourth technical contribution. It "
        "provides evidence about whether the implemented contributions behave as intended, where "
        "they fail, and what operational costs they introduce."
    )
    add_paragraph(
        doc,
        "The project contains eleven internal evaluation areas. The thesis emphasizes Evaluation 3 "
        "(task availability), Evaluation 7 (AI explanation quality), Evaluation 10 (coverage and "
        "human utility), and Evaluation 11 (system performance). Evaluations 1, 2, 4, 5, and 6 are "
        "supporting checks for import, mapping, SQL, API, and visualization behavior. Evaluations 8 "
        "and 9 are not presented as completed evidence and remain future work."
    )
    add_heading(doc, "5.1.1 Evaluation Questions and Evidence", 3)
    add_table(
        doc,
        ["Evaluation", "Question", "Main evidence"],
        [
            ["Supporting checks", "Is the data and execution path reliable enough for primary evaluation?", "Import/mapping logs and deterministic implementation checks"],
            ["3", "Does one fixed task scope produce defensible dataset-specific availability?", "52-task UCI and OULAD availability records"],
            ["7", "How do baseline and task-aware AI evidence modes compare?", "104 explanation records per dataset, judge outputs, and paired scores"],
            ["10", "What analytical coverage and user-facing transparency are provided?", "Coverage ratios, warnings, diagnostics, and qualitative review"],
            ["11", "What latency, stability, resource, token, cost, and import behavior is observed?", "Raw performance logs and consolidated benchmark report"],
        ],
        [1.2, 3.0, 2.3],
    )
    add_heading(doc, "5.1.2 Datasets", 3)
    add_paragraph(
        doc,
        "The UCI Portuguese student-performance dataset is a compact offline-oriented table with "
        "649 students, demographic and family attributes, support and lifestyle variables, absences, "
        "and three grade points. It has no learning-platform clickstream or submission timestamps. "
        "OULAD is a multi-file online-learning dataset containing modules, presentations, "
        "registrations, assessments, submissions, resources, and more than ten million raw student-"
        "VLE interaction rows. The contrast makes task-availability differences observable."
    )
    add_table(
        doc,
        ["Property", "UCI Portuguese", "OULAD imported batch"],
        [
            ["Students", "649", "28,785"],
            ["Assessments/results", "3 / 1,947", "206 / 173,912"],
            ["Events/engagement", "0 / 0", "6,364 / 8,459,320"],
            ["Strong signals", "Grades, outcomes, demographics, absences, study time, family/lifestyle", "Grades, outcomes, submissions, socioeconomic context, clickstream, temporal activity"],
            ["Important absent signals", "Clickstream, submission timing, temporal activity, socioeconomic context", "UCI-style lifestyle, family, absence, and study-time fields"],
        ],
        [1.45, 2.5, 2.55],
    )
    add_heading(doc, "5.1.3 Frozen Scope and Environment", 3)
    add_paragraph(
        doc,
        "The implementation registry contains 57 records at the time of writing, while the reported "
        "evaluation freezes a public 52-task scope. For AI evaluation, every task has a baseline and "
        "a task-aware explanation, producing 52 pairs and 104 records per dataset, or 208 target "
        "records in total. Freezing the scope prevents later experimental tasks from changing the "
        "reported denominator."
    )
    add_table(
        doc,
        ["Environment item", "Configuration"],
        [
            ["Machine", "Windows 10 x64; Intel Core i5-1135G7; 8 logical processors; approximately 8 GB RAM"],
            ["Runtime", "Node.js v22.14.0, local Express backend, PostgreSQL, local Python AI service"],
            ["Backend runs", "5 warm-up and 30 measured sequential runs per scenario"],
            ["AI runs", "2 warm-up and 10 measured runs per dataset"],
            ["Import runs", "2 warm-up and 10 measured UCI imports; one dedicated full OULAD import"],
            ["Timeout/percentile", "120,000 ms request timeout; nearest-rank percentiles"],
        ],
        [2.0, 4.5],
    )
    add_source_note(doc, "Source: Docs/evaluation_logs/system_performance/SYSTEM_PERFORMANCE_EVALUATION.md.")

    add_heading(doc, "5.2 Supporting System Verification", 2)
    add_paragraph(
        doc,
        "Primary results depend on earlier stages. A wrong mapping can create a false capability, "
        "and an incorrect SQL result can make a grounded explanation faithfully describe the wrong "
        "number. Supporting evidence is therefore reported first. It establishes an adequate "
        "evaluation foundation but is not claimed as formal proof of every software path."
    )
    add_heading(doc, "5.2.1 Profiling and Mapping", 3)
    add_paragraph(
        doc,
        "The latest UCI profiling record detected the semicolon delimiter with confidence 0.95, "
        "identified 649 rows and 33 columns, and resolved a request initially labeled CUSTOM to UCI. "
        "Thirty-one columns received high-confidence mappings, two fields were intentionally ignored, "
        "and no field remained in needs-review. The three grade columns were recognized as repeated "
        "score sources for later expansion into assessment records."
    )
    add_table(
        doc,
        ["Profiling/mapping measure", "Result"],
        [
            ["Rows / columns", "649 / 33"],
            ["Delimiter", "Semicolon; confidence 0.95"],
            ["Detected role", "UCI student_course_assessment_bundle"],
            ["High-confidence mapped fields", "31"],
            ["Intentionally unmapped", "2: reason and nursery"],
            ["Needs review", "0"],
            ["System-derived", "source_dataset from upload metadata"],
        ],
        [3.3, 3.2],
    )
    add_source_note(doc, "Source: profiling_mapping_20260620T075814Z_04a8bc.json.")
    add_heading(doc, "5.2.2 Import and Canonical Conversion", 3)
    ic = evidence["import_conversion"]
    add_paragraph(
        doc,
        "The confirmed UCI mapping was executed as a real database import. The 649 flat source rows "
        "were decomposed into 649 students, one course, one class, 649 enrollments, three assessments, "
        "and 1,947 assessment results. Repeated candidate course, class, and assessment entities were "
        "deduplicated before persistence. The import completed successfully in 3,322 ms."
    )
    add_table(
        doc,
        ["Canonical table", "Inserted rows"],
        [[key, str(value)] for key, value in ic["canonical_rows"].items()],
        [3.3, 2.7],
    )
    add_paragraph(
        doc,
        "Post-import validation found zero foreign-key errors and zero required-null or range errors. "
        "Feature checks found zero errors for pass_flag, week_of_class, registration_lead_time, "
        "log_click_score, and student composite features. For capabilities absent in UCI, zero errors "
        "means no inconsistent populated value was found; it does not mean the source supplied that "
        "capability."
    )
    add_table(
        doc,
        ["Check", "Errors"],
        [
            ["Foreign keys", "0"], ["Required nulls/ranges", "0"], ["pass_flag", "0"],
            ["week_of_class", "0"], ["registration_lead_time", "0"],
            ["log_click_score", "0"], ["Student composite features", "0"],
        ],
        [4.1, 1.9],
    )
    add_source_note(doc, "Source: import_conversion_20260620T080702Z_043fc1.json.")
    add_heading(doc, "5.2.3 SQL, API, and Visualization Checks", 3)
    add_paragraph(
        doc,
        "The analytics path checks capability before execution, binds server-owned SQL parameters, "
        "normalizes results by query_labels, and validates required output columns. Frontend tests "
        "cover supported adapters, semantic-field mapping, real-zero preservation, null filtering, "
        "risk cards, checklists, and diagnostics. These checks support later use of analytical "
        "outputs, but no consolidated numeric report currently proves every SQL task and chart; a "
        "universal correctness claim would therefore be too strong."
    )

    add_heading(doc, "5.3 Evaluation 3: Task Availability", 2)
    add_paragraph(
        doc,
        "The availability runner evaluates the same 52 tasks for a specific batch and class, storing "
        "the capability snapshot, final status, layer results, matched and missing requirements, "
        "warnings, confidence, and UI clickability. In the evaluated interface, only executable "
        "tasks are clickable."
    )
    add_heading(doc, "5.3.1 Capability Evidence", 3)
    add_table(
        doc,
        ["Evidence", "UCI Portuguese", "OULAD evaluated class"],
        [
            ["Enrollments/students", "649 / 649", "2,498 / 2,498"],
            ["Assessment results", "1,947", "11,451; 11,445 scored"],
            ["Assessment weeks", "3", "8"],
            ["Engagement", "0", "515,832 positive rows across 42 weeks"],
            ["Submission timestamps", "Unavailable", "11,451 populated rows"],
            ["Context strengths", "Family, lifestyle, absence, study time", "Socioeconomic, registration, clickstream"],
        ],
        [1.8, 2.2, 2.5],
    )
    add_heading(doc, "5.3.2 Status Results", 3)
    ta = evidence["task_availability"]
    add_table(
        doc,
        ["Dataset", "Total", "Executable", "Partial", "Insufficient", "Unsupported", "Disabled"],
        [
            ["UCI Portuguese", str(ta["UCI Portuguese"]["total"]), str(ta["UCI Portuguese"]["executable"]), str(ta["UCI Portuguese"]["partial"]), str(ta["UCI Portuguese"]["insufficient_data"]), str(ta["UCI Portuguese"]["unsupported"]), "28"],
            ["OULAD", str(ta["OULAD"]["total"]), str(ta["OULAD"]["executable"]), str(ta["OULAD"]["partial"]), str(ta["OULAD"]["insufficient_data"]), str(ta["OULAD"]["unsupported"]), "8"],
        ],
        [1.4, 0.7, 0.9, 0.8, 1.0, 0.9, 0.8],
    )
    add_source_note(doc, "Sources: task_availability_20260619T042411Z_265c90.json and task_availability_20260619T042455Z_0ef6b3.json.")
    add_paragraph(
        doc,
        "UCI exposes 24 executable tasks (46.15%), six partial tasks (11.54%), and 22 "
        "insufficient-data tasks (42.31%). OULAD exposes 44 executable tasks (84.62%) and eight "
        "partial tasks (15.38%). Neither dataset has an unsupported structural task. Differences "
        "therefore arise from populated semantic capabilities and sufficiency, not different task "
        "menus or absent deployed tables."
    )
    add_heading(doc, "5.3.3 Layer and Confidence Results", 3)
    add_table(
        doc,
        ["Measure", "UCI Portuguese", "OULAD"],
        [
            ["Structural failures", "0", "0"], ["Semantic failures", "28", "8"],
            ["Sufficiency failures", "22", "0"], ["Analytical warnings", "0", "0"],
            ["HIGH confidence", "30", "52"], ["LOW confidence", "22", "0"],
        ],
        [2.8, 1.85, 1.85],
    )
    add_paragraph(
        doc,
        "The UCI semantic-failure count exceeds the six final partial tasks because many semantic "
        "limitations also fail data sufficiency, whose final status takes precedence. All 22 "
        "insufficient UCI tasks receive LOW confidence. OULAD gives all tasks HIGH confidence but "
        "still marks eight partial where capabilities such as lifestyle or absence evidence are "
        "missing."
    )
    add_heading(doc, "5.3.4 Interpretation", 3)
    add_paragraph(
        doc,
        "Evaluation 3 supports the task-taxonomy contribution. One task scope remains stable across "
        "datasets while the executable subset responds to evidence. It also shows that schema "
        "compatibility alone is insufficient: the deployed canonical structure exists for both "
        "datasets, yet UCI correctly exposes fewer tasks because several domains are unpopulated."
    )
    add_paragraph(
        doc,
        "The result does not mean OULAD is intrinsically better. This registry contains many temporal, "
        "submission, engagement, and socioeconomic tasks. UCI contains family and lifestyle evidence "
        "that OULAD lacks. Availability measures compatibility with the declared taxonomy, not overall "
        "dataset quality."
    )

    add_heading(doc, "5.4 Evaluation 7: AI Explanation Quality", 2)
    add_paragraph(
        doc,
        "Evaluation 7 compares baseline_first_20_rows with task_aware_data_summarization. Both use "
        "completed SQL outputs and the same task context. A versioned judge scores faithfulness, "
        "numerical correctness, completeness, clarity, actionability, task relevance, and "
        "safety/fairness, with caps for major errors."
    )
    add_heading(doc, "5.4.1 Aggregate and Mode-Level Results", 3)
    add_table(
        doc,
        ["Dataset", "Target", "Scored", "Invalid", "Overall final avg.", "Major errors"],
        [
            ["UCI Portuguese", "104", "103", "1", "7.99", "2"],
            ["OULAD", "104", "104", "0", "7.92", "3"],
        ],
        [1.4, 0.8, 0.8, 0.8, 1.5, 1.0],
    )
    add_table(
        doc,
        ["Dataset", "Mode", "Scored", "Final average", "Major errors"],
        [
            ["UCI", "Baseline first 20 rows", "51", "7.96", "2"],
            ["UCI", "Task-aware summarization", "52", "8.02", "0"],
            ["OULAD", "Baseline first 20 rows", "52", "7.88", "1"],
            ["OULAD", "Task-aware summarization", "52", "7.97", "2"],
        ],
        [1.0, 2.5, 0.8, 1.2, 1.0],
    )
    add_paragraph(
        doc,
        "Task-aware mode has the higher average in both datasets. UCI task-aware records have no "
        "major errors, whereas OULAD task-aware mode has two major errors compared with one for the "
        "baseline. Broader evidence therefore changes the error profile but does not guarantee safer "
        "or more complete text."
    )
    add_heading(doc, "5.4.2 Paired Comparison", 3)
    add_table(
        doc,
        ["Dataset", "Comparable", "Task-aware wins", "Baseline wins", "Ties", "Mean delta"],
        [
            ["UCI Portuguese", "51", "25", "16", "10", "+0.09"],
            ["OULAD", "52", "18", "14", "20", "+0.09"],
        ],
        [1.4, 1.0, 1.2, 1.0, 0.7, 1.0],
    )
    add_paragraph(
        doc,
        "The paired improvement is positive but small. The UCI comparison excludes one invalid "
        "baseline record, so the paired mean is more appropriate than subtracting rounded mode "
        "averages. The winner counts and 20 OULAD ties reject a universal-superiority claim."
    )
    add_heading(doc, "5.4.3 Result-Size Analysis", 3)
    add_table(
        doc,
        ["Bucket", "Comparable", "Task-aware wins", "Baseline wins", "Ties"],
        [
            ["UCI <=20 rows", "45", "23", "14", "8"], ["UCI >20 rows", "6", "2", "2", "2"],
            ["OULAD <=20 rows", "39", "15", "10", "14"], ["OULAD >20 rows", "13", "3", "4", "6"],
        ],
        [1.75, 1.15, 1.25, 1.1, 0.8],
    )
    add_paragraph(
        doc,
        "When a result has at most 20 rows, the baseline already receives every row, so differences "
        "come from organization, completeness, numerical handling, relevance, actionability, and "
        "safety. For larger results, task-aware summaries can cover evidence beyond the baseline "
        "slice, but OULAD still contains four baseline wins, three task-aware wins, and six ties."
    )
    add_heading(doc, "5.4.4 Representative Cases", 3)
    add_paragraph(
        doc,
        "Large-result task-aware successes include UCI A-G02 (+1.35), UCI S-T15 (+3.11), OULAD "
        "A-G02 (+3.30), and OULAD S-T11 (+0.80). In A-G02, the baseline was capped for a major "
        "unsupported numerical claim. Counterexamples include UCI S-T09 (-0.55) and OULAD A-G14 "
        "(-0.80), where task-aware mode lost on correctness, completeness, safety, or relevance. "
        "OULAD A-B01 also gives baseline a 1.05 advantage after a task-aware missing-core-output error."
    )
    add_paragraph(
        doc,
        "Evaluation 7 therefore supports the AI contribution only in a constrained sense. The system "
        "produces structured and auditable explanations with high average scores, and task-aware "
        "preparation provides a modest average gain. It does not establish that the model is always "
        "correct or that one evidence strategy always produces a better educational insight."
    )
    add_source_note(doc, "Sources: official UCI and OULAD scoring reports and paired_mode_comparison aggregates under Docs/evaluation_v2/Runs/full_208/.")

    add_heading(doc, "5.5 Evaluation 10: Coverage and Human Utility", 2)
    add_paragraph(
        doc,
        "Evaluation 10 examines practical coverage and operational human utility. Coverage is "
        "quantified from the fixed task scope. Utility is reviewed through whether users can discover "
        "valid tasks, understand disabled states, inspect result-quality diagnostics, and request an "
        "explanation tied to evidence. This is not a controlled user study and does not measure "
        "learning gains, decision quality, or long-term adoption."
    )
    add_heading(doc, "5.5.1 Coverage", 3)
    add_table(
        doc,
        ["Dataset", "Executable", "Partial", "Insufficient", "Executable coverage"],
        [
            ["UCI Portuguese", "24", "6", "22", "46.15%"],
            ["OULAD", "44", "8", "0", "84.62%"],
        ],
        [1.45, 1.1, 0.9, 1.1, 1.45],
    )
    add_paragraph(
        doc,
        "Coverage is not maximized by pretending every task is valid. Correctly disabling 22 UCI "
        "tasks is useful evidence-aware behavior. OULAD's higher rate reflects richer online traces, "
        "while partial states preserve limited analytical reach without hiding missing enrichments."
    )
    add_heading(doc, "5.5.2 User-Facing Transparency", 3)
    add_table(
        doc,
        ["User need", "Visible evidence", "Operational value"],
        [
            ["Know if a task can run", "Status, badge, and clickable policy", "Avoids trial-and-error execution"],
            ["Understand unavailability", "Failed layer, missing capabilities, disabled reason, warnings", "Explains absence instead of appearing broken"],
            ["Inspect result quality", "Confidence and chart missing/skipped-row diagnostics", "Makes weak or filtered evidence visible"],
            ["Interpret output", "Structured AI insights with evidence and limitations", "Adds a reviewable narrative layer"],
            ["Continue without AI", "Degraded response while SQL result remains", "Separates analysis from optional explanation"],
        ],
        [1.55, 3.1, 1.85],
    )
    add_paragraph(
        doc,
        "The clearest utility is transparency about absence. With UCI, an engagement trend is disabled "
        "because engagement_tracking and temporal_activity are absent; with OULAD, the same task can "
        "become executable because those canonical capabilities are populated. The interface "
        "communicates consequences of evidence rather than relying only on dataset names."
    )
    add_heading(doc, "5.5.3 Limit of the Utility Claim", 3)
    add_paragraph(
        doc,
        "The evidence is formative and based on logs, interface behavior, explanation artifacts, and "
        "reviewer inspection. LLM-judge actionability is not equivalent to human usefulness. The "
        "defensible claim is limited to operational interpretability and evidence visibility. Future "
        "work should measure comprehension, trust calibration, task completion, and decision outcomes "
        "with real students and instructors."
    )

    add_heading(doc, "5.6 Evaluation 11: System Performance", 2)
    add_paragraph(
        doc,
        "Evaluation 11 measures latency, stability, throughput, resources, AI usage, and import "
        "performance. Recorded durations include backend processing, PostgreSQL, HTTP delivery, and "
        "external model latency for AI scenarios; they exclude browser rendering."
    )
    add_heading(doc, "5.6.1 Backend and Analytics", 3)
    add_table(
        doc,
        ["Scenario", "Dataset", "Runs", "Average ms", "P95 ms", "Maximum ms", "Error"],
        [
            ["Availability", "UCI", "30", "841.282", "1,249.040", "1,275.836", "0%"],
            ["Availability", "OULAD", "30", "28,790.822", "31,048.256", "91,670.026", "0%"],
            ["Simple A-B02", "UCI", "30", "16.761", "26.149", "37.832", "0%"],
            ["Simple A-B02", "OULAD", "30", "66.890", "102.880", "115.171", "0%"],
            ["Trend S-T01", "UCI", "30", "59.235", "71.140", "548.506", "0%"],
            ["Trend S-T01", "OULAD", "30", "107.205", "138.737", "139.987", "0%"],
        ],
        [1.25, 0.8, 0.6, 1.1, 0.95, 1.1, 0.55],
    )
    add_paragraph(
        doc,
        "Representative SQL analytics remain fast: simple aggregation averages 16.761 ms on UCI and "
        "66.890 ms on OULAD, while trend analysis averages 59.235 and 107.205 ms. OULAD availability "
        "is the main backend bottleneck at 28.79 seconds average, including one retained 91.67-second "
        "spike. Capability caching and consolidated count queries are required for frequent refresh."
    )
    add_heading(doc, "5.6.2 AI Latency and Cost", 3)
    add_table(
        doc,
        ["Dataset", "Runs", "Average ms", "P95 ms", "Tokens", "Cost", "Error"],
        [
            ["UCI", "10", "5,733.360", "6,284.419", "18,142", "USD 0.004909", "0%"],
            ["OULAD", "10", "6,075.344", "7,420.751", "21,084", "USD 0.005332", "0%"],
        ],
        [0.8, 0.7, 1.25, 1.1, 0.95, 1.25, 0.6],
    )
    add_paragraph(
        doc,
        "AI is much slower than SQL because it includes an external call. Across 20 measured calls, "
        "usage totals 39,226 tokens and estimated cost is USD 0.010241, with no measured failure or "
        "degradation. These results apply to S-T01 and gpt-4o-mini-2024-07-18; another model, prompt, "
        "or provider price would change them."
    )
    add_heading(doc, "5.6.3 Import Performance", 3)
    add_table(
        doc,
        ["Scenario", "Runs", "Scope", "Duration", "Rows/sec", "Error"],
        [
            ["UCI API import", "10", "6,490 source rows total", "2,408.664 ms average", "289.229", "0%"],
            ["OULAD full load", "1", "10,655,280 raw clickstream rows", "1,028,042.561 ms / 17.13 min", "10,364.629", "0%"],
        ],
        [1.3, 0.6, 2.0, 1.7, 0.85, 0.55],
    )
    add_paragraph(
        doc,
        "The repeated UCI import spends 132.779 ms on profiling, 100.316 ms on confirmation, and "
        "2,156.124 ms on import execution on average. The OULAD run converts 10,655,280 raw rows into "
        "8,459,320 engagement rows and 8,701,209 total canonical rows. Its engagement bulk phase "
        "takes approximately 13.88 of the total 17.13 minutes. This demonstrates thesis-scale "
        "feasibility, not production throughput."
    )
    add_heading(doc, "5.6.4 Stability and Resources", 3)
    add_table(
        doc,
        ["Scenario", "Avg CPU", "Peak CPU", "Avg RAM", "Peak RAM"],
        [
            ["UCI availability", "34.764%", "78.036%", "90.053%", "92.895%"],
            ["OULAD availability", "28.915%", "100.000%", "84.842%", "96.287%"],
            ["OULAD full import", "55.851%", "100.000%", "88.545%", "96.842%"],
        ],
        [2.1, 1.1, 1.1, 1.1, 1.1],
    )
    add_paragraph(
        doc,
        "All final scenarios completed without measured failures: 180 backend/analytics requests, "
        "20 AI calls, and 11 imports. CPU reached 100% and whole-OS RAM was high in OULAD scenarios. "
        "Because resource sampling includes development tools, services, and the database, these "
        "percentages cannot be assigned to one process. Stability is shown only for sequential local "
        "testing, not concurrent production load."
    )
    add_source_note(doc, "Source: Docs/evaluation_logs/system_performance/SYSTEM_PERFORMANCE_EVALUATION.md.")

    add_heading(doc, "5.7 Threats to Result Validity", 2)
    add_paragraph(
        doc,
        "Availability measures registry compatibility, not educational value. LLM-judge scores "
        "measure a project-specific rubric, not learning improvement. Coverage and transparency are "
        "proxies for utility. Performance uses selected scenarios and sequential local execution. "
        "These construct limits constrain the conclusions drawn from otherwise reproducible logs."
    )
    add_paragraph(
        doc,
        "Internal validity is strengthened by frozen scopes, paired AI comparisons, raw outputs, "
        "deterministic evidence, output-schema checks, and explicit invalid-record reporting. It is "
        "limited by model nondeterminism, one judge model, one invalid UCI record, and project-specific "
        "calibration. External validity is limited to two public datasets and one local machine."
    )

    add_heading(doc, "5.8 Summary of Experimental Findings", 2)
    add_paragraph(
        doc,
        "Supporting verification shows that the tested UCI source can be detected, mapped, normalized, "
        "and persisted with zero recorded integrity, range, and feature errors. Evaluation 3 provides "
        "the clearest evidence for task taxonomy: the same 52 tasks yield 24 executable tasks for UCI "
        "and 44 for OULAD because their populated capabilities differ."
    )
    add_paragraph(
        doc,
        "Evaluation 7 partially supports the AI contribution. Explanations score highly and task-aware "
        "preparation improves the paired score by +0.09 on both datasets, but baseline wins, ties, and "
        "task-aware major errors prevent a stronger claim. The defensible result is systematic, "
        "auditable evidence preparation rather than guaranteed superior insight."
    )
    add_paragraph(
        doc,
        "Evaluation 10 shows useful coverage and transparent unavailable-task handling, but not a "
        "completed large-scale user study. Evaluation 11 shows fast representative SQL analytics and "
        "stable sequential runs while exposing two practical bottlenecks: OULAD availability and "
        "external AI latency. These findings define the evidence boundary discussed in Chapter 6."
    )


def chapter_6(doc: Document):
    page_break(doc)
    add_heading(doc, "Chapter 6: DISCUSSION", 1)
    add_heading(doc, "6.1 Interpretation of the Results", 2)
    add_paragraph(
        doc,
        "This chapter interprets the experimental findings in relation to the four research questions "
        "and the three technical contributions. The central result is not that one universal dashboard "
        "can execute every analysis over every educational dataset. Instead, the results show that a "
        "shared architecture can represent heterogeneous data, expose a common vocabulary of analytical "
        "tasks, and determine which parts of that vocabulary are valid for a particular dataset. This "
        "distinction is important because structural interoperability does not imply analytical "
        "interchangeability. Two sources may be imported into the same canonical model while still "
        "supporting different questions because their observations, temporal coverage, assessment detail, "
        "and engagement signals differ."
    )
    add_paragraph(
        doc,
        "The findings also clarify the role of generative AI in the proposed system. The language model "
        "is not used to discover facts independently from the data or to replace deterministic analysis. "
        "It is positioned after task validation and SQL execution, where it translates completed analytical "
        "evidence into a structured explanation. The modest average improvement of the task-aware mode, "
        "together with cases in which the baseline performs equally well or better, shows that evidence "
        "preparation must itself be evaluated. More context is not automatically better context. The "
        "defensible contribution is therefore an inspectable explanation process, not a claim of universally "
        "superior AI-generated insight."
    )

    add_heading(doc, "6.2 Answers to the Research Questions", 2)
    add_heading(doc, "6.2.1 RQ1: Normalizing Heterogeneous Educational Data", 3)
    add_paragraph(
        doc,
        "RQ1 asks how heterogeneous educational datasets can be normalized into a reusable analytical "
        "structure. The implementation and supporting verification indicate that this can be achieved "
        "through a canonical relational schema combined with explicit source profiling, mapping, "
        "transformation, provenance, and integrity checks. The canonical entities provide stable grains "
        "for learners, courses, classes, enrollments, assessments, assessment results, events, and "
        "engagement. Source-specific fields are translated into these grains instead of being referenced "
        "directly by downstream tasks. As a result, task definitions and visualization contracts can be "
        "written against stable concepts rather than individual file layouts."
    )
    add_paragraph(
        doc,
        "The answer is qualified by the principle of non-fabrication. Canonicalization can reconcile "
        "names, identifiers, formats, and entity relationships, but it cannot create temporal activity, "
        "assessment attempts, or demographic attributes that the source never recorded. The schema is "
        "therefore reusable because it allows partial population and explicit absence, not because it "
        "forces all sources into an artificially complete representation. This design preserves the "
        "difference between a missing value inside a supported concept and an unsupported concept that "
        "has no evidential basis in the dataset."
    )

    add_heading(doc, "6.2.2 RQ2: Determining Meaningful Task Availability", 3)
    add_paragraph(
        doc,
        "RQ2 asks how a task taxonomy can determine which analyses are meaningful for each dataset. "
        "Evaluation 3 provides the strongest answer. The same fixed registry of 52 public tasks produced "
        "24 executable tasks for the UCI case and 44 executable tasks for OULAD. Because the registry was "
        "held constant, this difference is attributable to dataset capability rather than to separate "
        "dashboard implementations. The result demonstrates that task availability must be computed from "
        "the active evidence rather than assumed from the presence of an import record or a user-interface "
        "component."
    )
    add_paragraph(
        doc,
        "The four validation layers make this decision interpretable. Structural checks establish whether "
        "required canonical tables and relationships exist. Semantic checks determine whether necessary "
        "concepts and roles were mapped. Analytical checks examine whether suitable dimensions and measures "
        "are present. Quantitative checks assess whether the available observations are sufficient for the "
        "task contract. The resulting states distinguish executable, partial, insufficient-data, and "
        "unsupported cases. This is more informative than a binary feature flag because it communicates both "
        "the decision and its cause."
    )
    add_paragraph(
        doc,
        "Meaningfulness in this thesis is operational rather than pedagogically universal. An executable "
        "status means that the declared task requirements are satisfied by the available data and that the "
        "system can produce a contract-compliant result. It does not prove that the task is educationally "
        "valuable for every learner or institutional context. That stronger claim would require stakeholder "
        "research and outcome evaluation beyond the current scope."
    )

    add_heading(doc, "6.2.3 RQ3: Effectiveness of Evidence-Grounded AI Explanation", 3)
    add_paragraph(
        doc,
        "RQ3 asks how effectively task-aware AI can generate evidence-grounded explanations. Evaluation 7 "
        "shows that both evidence modes can produce generally strong explanations under the project rubric. "
        "The task-aware mode improved the paired score by approximately 0.09 on average for both UCI and "
        "OULAD. This repeated direction across two contrasting datasets suggests that structured evidence "
        "preparation can improve the use of analytical context. However, the magnitude is small, baseline "
        "wins and ties remain, and major errors are not eliminated. The result supports a limited claim of "
        "incremental improvement rather than categorical superiority."
    )
    add_paragraph(
        doc,
        "The paired design is important for interpretation because it compares explanations generated from "
        "the same analytical task and underlying result. Differences can therefore be linked more directly "
        "to evidence preparation than in an unpaired aggregate comparison. At the same time, LLM generation "
        "and LLM-based judging are nondeterministic and model-dependent. The scores should be read as evidence "
        "from a controlled project protocol, not as an absolute measurement of explanation truth or human "
        "learning benefit."
    )
    add_paragraph(
        doc,
        "The strongest aspect of the AI design is traceability. Task identifiers, result datasets, "
        "deterministic summaries, prompts, model outputs, judge inputs, rubric scores, and error records can "
        "be inspected as a chain. This makes unsupported claims and omissions easier to identify than in an "
        "interface that displays an unlogged model response. The remaining weakness is that traceability does "
        "not itself prevent error. Human review and conservative wording remain necessary, especially when "
        "an explanation concerns learner risk or recommends an intervention."
    )

    add_heading(doc, "6.2.4 RQ4: Coverage, Utility, and Performance", 3)
    add_paragraph(
        doc,
        "RQ4 asks what the coverage, human-utility, and performance evaluations reveal about the complete "
        "system. Evaluation 10 indicates that the architecture covers a useful range of student-facing and "
        "administrative analytical intents while making unavailable tasks visible through statuses, reasons, "
        "and warnings. This supports operational transparency: users can distinguish a valid zero or empty "
        "result from a task that cannot be supported by the imported data. The evidence does not constitute "
        "a large-scale usability study, so the term human utility is restricted to inspectability, clarity of "
        "system feedback, and reviewer-observed usefulness."
    )
    add_paragraph(
        doc,
        "Evaluation 11 shows that representative deterministic analytics are fast enough for interactive "
        "thesis-scale use and that the import pipeline can process the large OULAD activity source. It also "
        "identifies two practical bottlenecks. OULAD availability analysis requires more work because its "
        "capability profile and task evidence are larger, while AI response time is dominated by the external "
        "model call rather than local SQL execution. These findings suggest that the architecture is viable as "
        "a prototype but still requires caching, incremental validation, concurrency testing, and deployment "
        "monitoring before production use."
    )

    add_heading(doc, "6.3 Discussion of the Three Contributions", 2)
    add_heading(doc, "6.3.1 Contribution 1: Canonical Educational Schema", 3)
    add_paragraph(
        doc,
        "The canonical schema contributes a stable analytical boundary between heterogeneous sources and "
        "downstream services. Its value is not limited to reducing column-name differences. By defining entity "
        "grain, keys, relationships, provenance, and feature ownership, the schema constrains how imported "
        "records may be interpreted. Analytics can refer to enrollment, assessment result, or engagement as "
        "explicit concepts rather than infer those concepts repeatedly from source-specific columns. This "
        "reduces duplicated assumptions across SQL, APIs, charts, and AI prompts."
    )
    add_paragraph(
        doc,
        "The UCI and OULAD cases illustrate complementary strengths. UCI provides compact learner, school, "
        "family, behavior, and grade attributes, whereas OULAD provides richer course presentation, assessment, "
        "registration, and activity records. A rigid intersection schema would discard much of this information, "
        "while a union without semantics would merely collect unrelated fields. The proposed model uses shared "
        "core entities with optional population and provenance, allowing common analysis where concepts align "
        "and explicit limitation where they do not."
    )
    add_paragraph(
        doc,
        "The main cost is mapping effort. Ambiguous source columns require human confirmation, and new "
        "institutional sources may require schema extensions or transformation rules. The schema should "
        "therefore be treated as an evolving contract with versioning and migration support, not as a final "
        "universal ontology. Its contribution is a practical thesis-scale method for disciplined normalization."
    )

    add_heading(doc, "6.3.2 Contribution 2: Task Taxonomy and Availability Mechanism", 3)
    add_paragraph(
        doc,
        "The second contribution shifts the unit of dashboard design from the page or chart to the analytical "
        "task. Each task combines audience, intent, requirements, parameters, SQL behavior, output shape, "
        "visualization metadata, and explanation strategy. This creates a vocabulary that can be inspected "
        "independently from the interface and reused by the backend, frontend, and evaluation runners. The "
        "registry thereby serves as both executable configuration and a testable analytical specification."
    )
    add_paragraph(
        doc,
        "The availability mechanism is essential because metadata alone can reproduce hidden assumptions if "
        "requirements are incomplete. The observed difference between UCI and OULAD availability validates "
        "the need for dataset-sensitive checks, but it also reveals a maintenance obligation: requirement "
        "metadata, capability detection, and SQL contracts must evolve together. A newly introduced task can "
        "be incorrectly enabled if its requirements omit a necessary signal. Registry review and regression "
        "tests are therefore part of the method, not optional implementation hygiene."
    )
    add_paragraph(
        doc,
        "Compared with a hardcoded dashboard, the proposed approach requires more explicit metadata and an "
        "additional validation stage. For a single stable dataset, hardcoding may be faster initially. The "
        "metadata-driven approach becomes more valuable when datasets change, when multiple sources are "
        "compared, or when users need reasons for disabled analysis. Its trade-off is increased up-front "
        "specification in exchange for reuse, transparency, and more systematic testing."
    )

    add_heading(doc, "6.3.3 Contribution 3: Evidence-Grounded AI Explanation", 3)
    add_paragraph(
        doc,
        "The third contribution separates numerical computation from narrative interpretation. SQL and "
        "deterministic transformation produce the analytical result, while the language model receives a "
        "bounded representation of that result together with task context and limitations. This separation "
        "reduces the need for the model to infer numbers, recover task meaning from raw rows, or invent a "
        "visual interpretation without metadata. Structured response fields also allow the interface and "
        "evaluation pipeline to distinguish summaries, observations, recommendations, warnings, and evidence."
    )
    add_paragraph(
        doc,
        "The Evaluation 7 results show why the contribution should be framed around grounding rather than "
        "intelligence. Task-aware evidence can improve completeness and relevance, particularly when raw "
        "results are large or multi-part, but summarization can omit details or overemphasize selected patterns. "
        "The baseline may be competitive when the result is already compact. An effective system should "
        "therefore select an evidence strategy according to result shape and retain a fallback path rather than "
        "apply one elaborate summarizer to every task."
    )
    add_paragraph(
        doc,
        "Educational explanations also carry ethical risk. Descriptive patterns can be mistaken for causes, "
        "group averages can be applied to individuals, and recommendations can acquire unwarranted authority "
        "when written fluently. The system addresses these concerns through task-level safety instructions, "
        "warnings, confidence information, and explicit evidence boundaries. These controls reduce risk but do "
        "not justify autonomous decisions about students. Final interpretation remains a human responsibility."
    )

    add_heading(doc, "6.4 Cross-Cutting Architectural Trade-offs", 2)
    add_heading(doc, "6.4.1 Reuse versus Source Fidelity", 3)
    add_paragraph(
        doc,
        "A shared model promotes reuse, but excessive abstraction can erase source-specific meaning. The "
        "architecture addresses this tension by preserving source provenance and permitting optional canonical "
        "coverage. Common fields support reusable tasks, while unmapped or source-specific information remains "
        "available for diagnosis and future extension. The trade-off should be managed through documented "
        "mapping decisions and schema evolution rather than by forcing every source into identical completeness."
    )

    add_heading(doc, "6.4.2 Determinism versus Generative Flexibility", 3)
    add_paragraph(
        doc,
        "Deterministic SQL provides repeatable measures and clear failure modes, whereas generative models "
        "provide flexible language and audience-sensitive interpretation. Combining them is useful only when "
        "their responsibilities remain separate. If the language model computes metrics, reproducibility is "
        "weakened. If deterministic templates are required to express every insight, communication becomes "
        "rigid. The proposed sequence uses deterministic computation as the evidential core and generative "
        "language as a constrained communication layer."
    )

    add_heading(doc, "6.4.3 Transparency versus Interface Simplicity", 3)
    add_paragraph(
        doc,
        "Availability reasons, confidence, provenance, and AI warnings improve transparency but can burden the "
        "interface. The system must present enough diagnostic information to prevent misleading interpretation "
        "without overwhelming ordinary users. A practical design is progressive disclosure: show a concise "
        "status and explanation by default, while preserving detailed capability, evidence, and log information "
        "for reviewers and administrators."
    )

    add_heading(doc, "6.5 Practical Implications", 2)
    add_heading(doc, "6.5.1 Implications for Learners and Educators", 3)
    add_paragraph(
        doc,
        "For learners and educators, the primary benefit is not the number of charts but the clarity of the "
        "analytical contract. A valid task can present a deterministic result and a bounded explanation. An "
        "invalid task can state which concept or quantity is missing instead of displaying an empty chart or "
        "silently disappearing. This distinction supports more informed use of educational data and reduces "
        "the risk that absence of evidence is interpreted as evidence of absence."
    )
    add_paragraph(
        doc,
        "AI-generated recommendations should be treated as prompts for reflection, not instructions. Educators "
        "can use the explanation to identify patterns worth investigating, but they must consider classroom "
        "context, assessment design, learner circumstances, and data quality before acting. Learners should be "
        "able to understand what data contributed to an explanation and to challenge an interpretation that "
        "does not match their situation."
    )

    add_heading(doc, "6.5.2 Implications for System Designers", 3)
    add_paragraph(
        doc,
        "For system designers, the results support treating task validity as a first-class service. Import "
        "success, API success, and chart rendering are insufficient indicators of analytical correctness. "
        "Capability snapshots and requirement metadata should be tested together with SQL and output contracts. "
        "The evaluation logs further show the value of preserving intermediate artifacts so that a failure can "
        "be traced to mapping, availability, execution, evidence construction, generation, or judging."
    )
    add_paragraph(
        doc,
        "The performance results suggest clear optimization priorities. Deterministic task execution should "
        "remain lightweight, availability checks should reuse cached capability summaries where safe, and AI "
        "calls should use timeout, retry, degradation, and cost controls. Production deployment would also need "
        "access control, audit retention, privacy governance, model versioning, and monitoring beyond the scope "
        "of the current prototype."
    )

    add_heading(doc, "6.6 Reproducibility and Evidence Boundaries", 2)
    add_paragraph(
        doc,
        "The evaluation protocol records import outputs, task availability decisions, analytics results, model "
        "invocations, judge responses, aggregate scores, paired comparisons, and performance measurements. "
        "This artifact chain allows a reported result to be checked against its inputs and configuration. It "
        "also reduces the temptation to select only visually successful examples. Reproducibility here means "
        "that the procedure, scope, and evidence are inspectable and can be rerun under a recorded environment."
    )
    add_paragraph(
        doc,
        "Reproducibility does not remove all uncertainty. External models may change, repeated generations may "
        "differ, and local performance depends on hardware and system state. The thesis therefore reports "
        "results with their run context and avoids treating one score or latency value as universal. The "
        "evaluation protocol supports the three contributions, but it is not framed as an independent technical contribution."
    )

    add_heading(doc, "6.7 Threats to Validity", 2)
    add_heading(doc, "6.7.1 Construct Validity", 3)
    add_paragraph(
        doc,
        "Task availability operationalizes whether declared data requirements are satisfied; it does not measure "
        "educational importance. The AI rubric measures faithfulness, numerical correctness, completeness, "
        "relevance, clarity, actionability, and safety under a project-defined scoring procedure; it does not "
        "directly measure comprehension or learning improvement. Evaluation 10 uses coverage, visible states, "
        "warnings, and reviewer inspection as utility evidence, so its claim remains narrower than formal "
        "usability or user-impact evaluation."
    )

    add_heading(doc, "6.7.2 Internal Validity", 3)
    add_paragraph(
        doc,
        "Internal validity is strengthened by fixed task scopes, automated runners, deterministic SQL evidence, "
        "paired comparison, output-schema validation, and retention of raw artifacts. It is limited by the "
        "possibility of errors in ground-truth labels, requirement metadata, evidence summarization, and judge "
        "calibration. LLM nondeterminism and reliance on one principal judging configuration may affect scores. "
        "Manual inspection and invalid-record reporting reduce but do not eliminate these risks."
    )

    add_heading(doc, "6.7.3 External Validity", 3)
    add_paragraph(
        doc,
        "External validity is limited by the use of two public datasets and a thesis-scale deployment. UCI and "
        "OULAD are deliberately contrasting, but they do not represent every learning management system, "
        "institutional policy, data quality condition, language, or student population. Performance was measured "
        "on one local environment with sequential scenarios rather than production concurrency. Findings about "
        "reuse and feasibility are therefore transferable as architectural evidence, not as universal deployment "
        "guarantees."
    )

    add_heading(doc, "6.7.4 Conclusion Validity and Unfinished Evaluations", 3)
    add_paragraph(
        doc,
        "The thesis emphasizes Evaluation 3, Evaluation 7, Evaluation 10, and Evaluation 11 because they have "
        "the clearest completed artifacts for the central claims. Evaluations 1, 2, 4, 5, and 6 provide supporting "
        "verification of the pipeline. Evaluation 8 on statistical validity and Evaluation 9 on early-warning "
        "utility are not treated as completed evidence. Consequently, the thesis does not claim causal validity, "
        "predictive intervention effectiveness, or statistically generalizable improvement in learner outcomes."
    )

    add_heading(doc, "6.8 Chapter Summary", 2)
    add_paragraph(
        doc,
        "The discussion shows that the three contributions form a dependency chain. The canonical schema "
        "defines what evidence is represented. The task taxonomy and availability validator determine what can "
        "be analyzed. The AI layer communicates completed analysis within a bounded evidence context. Evaluation "
        "3 demonstrates dataset-sensitive task validity, Evaluation 7 supports a modest benefit from task-aware "
        "evidence preparation, Evaluation 10 supports operational transparency and coverage, and Evaluation 11 "
        "establishes thesis-scale feasibility while identifying performance bottlenecks. These results support "
        "the proposed architecture without extending the claims to universal interoperability, autonomous "
        "educational decision-making, or proven learning improvement."
    )


def chapter_7(doc: Document):
    page_break(doc)
    add_heading(doc, "Chapter 7: CONCLUSION AND FUTURE WORK", 1)
    add_heading(doc, "7.1 Thesis Summary", 2)
    add_paragraph(
        doc,
        "This thesis addressed a recurring problem in learning analytics: educational data can be imported "
        "and visualized without establishing whether the same analytical meaning survives across datasets or "
        "whether an AI-generated explanation is supported by the available evidence. To address this problem, "
        "the work designed, implemented, and evaluated Smart Learning Dashboard with AI-Assisted Data Analysis "
        "as a metadata-driven and evidence-aware architecture. The system connects source profiling and mapping, "
        "canonical persistence, analytical task metadata, task availability validation, deterministic SQL, "
        "metadata-driven visualization, and bounded AI explanation in one traceable pipeline."
    )
    add_paragraph(
        doc,
        "The architecture was studied using two contrasting educational datasets. UCI student performance "
        "provides a compact learner-level source with demographic, social, behavioral, and grade attributes. "
        "OULAD provides relational course, registration, assessment, and activity records at a substantially "
        "larger scale. Their differences made it possible to test whether the system preserves a common "
        "analytical vocabulary without pretending that both sources support identical tasks."
    )

    add_heading(doc, "7.2 Summary of the Three Contributions", 2)
    add_heading(doc, "7.2.1 Canonical Educational Schema", 3)
    add_paragraph(
        doc,
        "The first contribution is a canonical educational schema with explicit entity grains, relationships, "
        "provenance, and mapping rules. It separates source-specific representation from downstream analysis "
        "and supports partial coverage without fabricating absent educational concepts. The schema enables "
        "shared task definitions to operate over stable entities while preserving enough provenance to inspect "
        "how each canonical record was produced."
    )

    add_heading(doc, "7.2.2 Task Taxonomy and Availability Mechanism", 3)
    add_paragraph(
        doc,
        "The second contribution is a task taxonomy operationalized through a metadata-driven registry and a "
        "four-layer availability validator. Tasks declare their audience, intent, parameters, required "
        "capabilities, SQL behavior, output contract, visualization metadata, and explanation strategy. The "
        "validator evaluates structural, semantic, analytical, and quantitative readiness before ordinary "
        "execution. This converts hidden assumptions into explicit, testable requirements and allows the "
        "interface to communicate why an analysis is executable, partial, insufficient, or unsupported."
    )

    add_heading(doc, "7.2.3 Evidence-Grounded AI Explanation", 3)
    add_paragraph(
        doc,
        "The third contribution is an AI explanation pipeline that operates on completed analytical evidence. "
        "The model receives task context, named results, deterministic summaries, limitations, and safety "
        "instructions rather than unrestricted access to raw institutional data or responsibility for numerical "
        "computation. Structured outputs and retained evaluation artifacts make explanations inspectable. The "
        "contribution is evidence discipline and traceability, not a new language model and not a guarantee that "
        "every generated explanation is correct."
    )
    add_paragraph(
        doc,
        "These contributions are connected rather than independent. The canonical schema defines the available "
        "evidence, the task taxonomy determines valid analytical use of that evidence, and the AI layer "
        "communicates the completed result. The evaluation protocol provides supporting evidence for this chain "
        "and is not presented as an additional technical contribution."
    )

    add_heading(doc, "7.3 Main Findings", 2)
    add_paragraph(
        doc,
        "For RQ1, the import and integrity evidence indicates that heterogeneous data can be normalized into a "
        "shared analytical structure when mapping, transformation, provenance, and partial coverage are explicit. "
        "For RQ2, Evaluation 3 shows that a fixed taxonomy can yield different valid analytical spaces: 24 of 52 "
        "public tasks were executable for UCI, compared with 44 of 52 for OULAD. This confirms that successful "
        "import should not be equated with universal task support."
    )
    add_paragraph(
        doc,
        "For RQ3, Evaluation 7 shows generally strong explanation quality and an average paired improvement of "
        "approximately +0.09 for task-aware evidence on both datasets. The effect is modest, baseline explanations "
        "sometimes win, and major errors remain possible. The result supports structured evidence preparation as "
        "a useful mechanism while rejecting the stronger claim that task-aware prompting is always superior."
    )
    add_paragraph(
        doc,
        "For RQ4, Evaluation 10 supports analytical coverage and transparent handling of unavailable tasks, but it "
        "does not prove learning improvement or complete usability. Evaluation 11 shows that representative SQL "
        "analytics and import operations are feasible at thesis scale. It also shows that OULAD availability "
        "checking and external AI calls are the main optimization targets. Together, the evaluations establish a "
        "working and auditable prototype while defining the boundary of its present claims."
    )

    add_heading(doc, "7.4 Limitations", 2)
    add_paragraph(
        doc,
        "The study has several limitations. First, the empirical scope is restricted to two public datasets. "
        "They differ meaningfully in structure and scale, but they cannot represent the full diversity of "
        "institutional information systems, data governance practices, languages, curricula, and learner groups. "
        "Second, the canonical schema and task registry are thesis-scale artifacts. New sources and analytical "
        "domains may require schema evolution, additional capability rules, and revised task contracts."
    )
    add_paragraph(
        doc,
        "Third, the explanation evaluation depends on a project-specific rubric and an LLM-based judge. Paired "
        "comparison and artifact retention strengthen the protocol, but judge bias, model drift, prompt sensitivity, "
        "and generation nondeterminism remain. Fourth, Evaluation 10 provides operational and reviewer-based "
        "evidence rather than a controlled study with a large number of learners and educators. Claims about "
        "human utility are therefore limited to transparency, interpretability, and plausible workflow support."
    )
    add_paragraph(
        doc,
        "Finally, the system was evaluated in a local sequential environment. Production deployment would introduce "
        "concurrency, authentication, privacy, retention, failure recovery, cost, and monitoring requirements that "
        "were not tested fully. Evaluation 8 on statistical validity and Evaluation 9 on early-warning utility are "
        "not treated as completed evidence. The thesis consequently makes no claim of causal inference, validated "
        "risk prediction, autonomous intervention, or demonstrated improvement in learning outcomes."
    )

    add_heading(doc, "7.5 Future Work", 2)
    add_heading(doc, "7.5.1 Broader Data and Schema Validation", 3)
    add_paragraph(
        doc,
        "Future work should evaluate the canonical schema with additional learning management systems, student "
        "information systems, assessment platforms, and event standards. This would reveal which entities and "
        "mapping rules generalize and which require extension. Schema versions, migration procedures, mapping "
        "provenance, and automated compatibility tests should be introduced so that new concepts can be added "
        "without silently changing existing analytical meaning."
    )

    add_heading(doc, "7.5.2 Stronger Task and Statistical Validation", 3)
    add_paragraph(
        doc,
        "The task registry should be expanded with machine-checkable contracts for minimum sample size, temporal "
        "coverage, grouping validity, missingness, and statistical assumptions. Evaluation 8 should be completed "
        "with deterministic reference calculations, appropriate uncertainty reporting, and tests for tasks that "
        "use comparisons or statistical summaries. Availability rules should also be validated by independent "
        "reviewers to estimate agreement and identify underspecified task requirements."
    )

    add_heading(doc, "7.5.3 Early-Warning and Intervention Evaluation", 3)
    add_paragraph(
        doc,
        "Evaluation 9 requires a separate longitudinal design. Future work should define prediction times, outcomes, "
        "cohort splits, leakage controls, calibration measures, fairness checks, and intervention protocols before "
        "presenting early-warning claims. Predictive performance alone is insufficient. A useful early-warning "
        "system must also establish whether educators can act on alerts, whether interventions are timely, and "
        "whether harms from false positives and false negatives are acceptable."
    )

    add_heading(doc, "7.5.4 Human-Centered Evaluation", 3)
    add_paragraph(
        doc,
        "A formal study with learners, instructors, and administrators should examine comprehension, trust, task "
        "completion, diagnostic usefulness, and the effect of availability explanations. AI explanations should be "
        "compared with deterministic templates and no-explanation controls. Interviews and think-aloud sessions can "
        "identify whether evidence references and warnings are understandable or whether they create unnecessary "
        "cognitive load. Such work is necessary before making stronger claims about educational utility."
    )

    add_heading(doc, "7.5.5 AI Reliability and Governance", 3)
    add_paragraph(
        doc,
        "The explanation service can be improved through adaptive evidence selection, numerical claim verification, "
        "citation to result cells or aggregates, and automatic detection of unsupported causal language. Multiple "
        "model and judge configurations should be evaluated to measure robustness. Production governance should "
        "record model versions, prompts, costs, safety events, and human overrides, while enforcing role-based access "
        "and data minimization for sensitive educational information."
    )

    add_heading(doc, "7.5.6 Performance and Deployment", 3)
    add_paragraph(
        doc,
        "Future engineering work should cache stable capability snapshots, evaluate incremental availability checks, "
        "and profile large OULAD-style sources under concurrent load. AI calls require timeout, retry, rate-limit, "
        "fallback, and budget policies. End-to-end deployment tests should include authentication, authorization, "
        "privacy controls, backup and recovery, observability, and failure injection. These steps would move the "
        "artifact from a validated prototype toward an institution-ready service."
    )

    add_heading(doc, "7.6 Concluding Statement", 2)
    add_paragraph(
        doc,
        "The thesis demonstrates that a learning analytics dashboard can be designed around explicit evidence "
        "contracts rather than a fixed dataset and a collection of isolated charts. Canonical representation makes "
        "heterogeneous data analyzable, task metadata makes analytical validity inspectable, and grounded AI makes "
        "narrative interpretation auditable. The evaluation also shows the limits of this approach: available tasks "
        "remain dataset-dependent, task-aware AI offers only a modest average improvement, and prototype feasibility "
        "does not establish institutional or educational impact. Within these boundaries, the proposed architecture "
        "provides a coherent foundation for learning analytics systems that are more reusable, transparent, and "
        "responsible in how they transform data into analysis and explanation."
    )

    add_heading(doc, "References", 1)
    add_heading(doc, "Academic References", 2)
    academic_references = [
        "1EdTech Consortium. (n.d.). Caliper Analytics Specification.",
        "Advanced Distributed Learning Initiative. (n.d.). Experience API (xAPI) Specification.",
        "Baker, R. S., & Inventado, P. S. (2014). Educational data mining and learning analytics. In J. A. Larusson & B. White (Eds.), Learning Analytics: From Research to Practice. Springer. https://doi.org/10.1007/978-1-4614-3305-7_4",
        "Bodily, R., & Verbert, K. (2017). Review of research on student-facing learning analytics dashboards and educational recommender systems. IEEE Transactions on Learning Technologies, 10(4), 405-418. https://doi.org/10.1109/TLT.2017.2740172",
        "Brehmer, M., & Munzner, T. (2013). A multi-level typology of abstract visualization tasks. IEEE Transactions on Visualization and Computer Graphics, 19(12), 2376-2385. https://doi.org/10.1109/TVCG.2013.124",
        "Cortez, P., & Silva, A. M. G. (2008). Using data mining to predict secondary school student performance. Proceedings of the 5th Future Business Technology Conference, 5-12.",
        "Ji, Z., Lee, N., Frieske, R., Yu, T., Su, D., Xu, Y., Ishii, E., Bang, Y. J., Madotto, A., & Fung, P. (2023). Survey of hallucination in natural language generation. ACM Computing Surveys, 55(12). https://doi.org/10.1145/3571730",
        "Kuzilek, J., Hlosta, M., & Zdrahal, Z. (2017). Open University Learning Analytics dataset. Scientific Data, 4, 170171. https://doi.org/10.1038/sdata.2017.171",
        "Lewis, P., Perez, E., Piktus, A., Petroni, F., Karpukhin, V., Goyal, N., Kuttler, H., Lewis, M., Yih, W.-t., Rocktaschel, T., Riedel, S., & Kiela, D. (2020). Retrieval-augmented generation for knowledge-intensive NLP tasks. Advances in Neural Information Processing Systems, 33, 9459-9474.",
        "Schwendimann, B. A., Rodriguez-Triana, M. J., Vozniuk, A., Prieto, L. P., Boroujeni, M. S., Holzer, A., Gillet, D., & Dillenbourg, P. (2017). Perceiving learning at a glance: A systematic literature review of learning dashboard research. IEEE Transactions on Learning Technologies, 10(1), 30-41. https://doi.org/10.1109/TLT.2016.2599522",
        "Siemens, G. (2013). Learning analytics: The emergence of a discipline. American Behavioral Scientist, 57(10), 1380-1400. https://doi.org/10.1177/0002764213498851",
        "Slade, S., & Prinsloo, P. (2013). Learning analytics: Ethical issues and dilemmas. American Behavioral Scientist, 57(10), 1510-1529. https://doi.org/10.1177/0002764213479366",
        "UNESCO. (2023). Guidance for generative AI in education and research. UNESCO.",
        "Verbert, K., Duval, E., Klerkx, J., Govaerts, S., & Santos, J. L. (2013). Learning analytics dashboard applications. American Behavioral Scientist, 57(10), 1500-1509. https://doi.org/10.1177/0002764213479363",
        "Zheng, L., Chiang, W.-L., Sheng, Y., Zhuang, S., Wu, Z., Zhuang, Y., Lin, Z., Li, Z., Li, D., Xing, E. P., Zhang, H., Gonzalez, J. E., & Stoica, I. (2023). Judging LLM-as-a-judge with MT-Bench and Chatbot Arena. Advances in Neural Information Processing Systems, 36.",
    ]
    for reference in academic_references:
        add_reference_entry(doc, reference)

    add_heading(doc, "Project Technical Artifacts", 2)
    project_references = [
        "Project documentation. Canonical schema specification. Docs/canonical_schema.md.",
        "Project documentation. System design. Docs/system_design.md.",
        "Project documentation. System architecture report. Docs/system-architecture-report.md.",
        "Project documentation. API contract and Swagger specification. Docs/06_API_Contract_Swagger.md.",
        "Project evaluation artifacts. Task availability logs. Docs/evaluation_logs/task_availability/.",
        "Project evaluation artifacts. Import and conversion logs. Docs/evaluation_logs/import_conversion/.",
        "Project evaluation artifacts. System performance evaluation. Docs/evaluation_logs/system_performance/SYSTEM_PERFORMANCE_EVALUATION.md.",
        "Project evaluation artifacts. Official AI explanation scoring reports. Docs/evaluation_v2/Runs/full_208/.",
    ]
    for reference in project_references:
        add_reference_entry(doc, reference)

    add_heading(doc, "Appendices", 1)
    add_heading(doc, "Appendix A: Canonical Schema and Mapping Specifications", 2)
    add_paragraph(doc, "Placeholder for detailed canonical table definitions, field mappings, and transformation rules.")
    add_heading(doc, "Appendix B: Full Task Taxonomy and Registry", 2)
    add_paragraph(doc, "Placeholder for the complete list of public and experimental tasks, required capabilities, and output contracts.")
    add_heading(doc, "Appendix C: Task Availability Ground Truth", 2)
    add_paragraph(doc, "Placeholder for task availability matrices and reviewer notes.")
    add_heading(doc, "Appendix D: AI Explanation Rubric and Examples", 2)
    add_paragraph(doc, "Placeholder for judge rubric, scoring schema, selected examples, and paired comparison cases.")
    add_heading(doc, "Appendix E: Coverage and Human-Utility Review Artifacts", 2)
    add_paragraph(doc, "Placeholder for coverage tables, disabled-task examples, and human review notes.")
    add_heading(doc, "Appendix F: Performance Logs and Reproducibility Instructions", 2)
    add_paragraph(doc, "Placeholder for benchmark commands, hardware environment, and links to raw JSON logs.")


def structural_audit(doc: Document) -> dict:
    headings = []
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        text = p.text.strip()
        if style.startswith("Heading") and text:
            headings.append({"style": style, "text": text})
    issues = []
    required = [
        "ABSTRACT",
        "Chapter 1: INTRODUCTION",
        "Chapter 2: BACKGROUND AND RELATED WORK",
        "Chapter 3: PROPOSED METHODOLOGY",
        "Chapter 4: SYSTEM DESIGN AND IMPLEMENTATION",
        "Chapter 5: EXPERIMENTS AND RESULTS",
        "Chapter 6: DISCUSSION",
        "Chapter 7: CONCLUSION AND FUTURE WORK",
    ]
    heading_texts = {h["text"] for h in headings}
    for item in required:
        if item not in heading_texts:
            issues.append(f"Missing required heading: {item}")
    all_text = "\n".join(p.text for p in doc.paragraphs)
    if "FUTIRE" in all_text:
        issues.append("Typo FUTIRE remains")
    if all_text.count("3.4.3 Runtime Derived Metrics") > 1 or "3.4.4 Runtime Derived Metrics" in all_text:
        issues.append("Duplicate Runtime Derived Metrics heading remains")
    return {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "headings": headings,
        "issues": issues,
    }


def main():
    if not BASE_DOCX.exists():
        raise FileNotFoundError(BASE_DOCX)
    evidence = collect_evidence()
    shutil.copyfile(BASE_DOCX, BUILD_DOCX)
    doc = Document(str(BUILD_DOCX))
    write_list_of_abbreviations(doc)
    remove_body_from_heading(doc, "ABSTRACT")

    # Start the rewritten thesis body after preserved front matter.
    page_break(doc)
    write_abstract(doc)
    chapter_1(doc)
    chapter_2(doc)
    chapter_3(doc)
    chapter_4(doc)
    chapter_5(doc, evidence)
    chapter_6(doc)
    chapter_7(doc)
    normalize_thesis_typography(doc)

    doc.save(str(BUILD_DOCX))
    shutil.copyfile(BUILD_DOCX, OUT_DOCX)
    BUILD_DOCX.unlink(missing_ok=True)
    audit = structural_audit(Document(str(OUT_DOCX)))
    audit_path = OUT_DOCX.with_suffix(".audit.json")
    audit_path.write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    if audit["issues"]:
        raise RuntimeError("Structural audit failed: " + "; ".join(audit["issues"]))
    print(json.dumps({"out": str(OUT_DOCX), "audit": str(audit_path), "paragraphs": audit["paragraphs"], "tables": audit["tables"], "headings": len(audit["headings"])}, indent=2))


if __name__ == "__main__":
    main()
